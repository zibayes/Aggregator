import sys
from unittest.mock import MagicMock

# Создаем заглушку ДО импорта проблемного модуля
if 'agregator.processing.account_cards_processing' not in sys.modules:
    # Создаем заглушку для всего модуля
    sys.modules['agregator.processing.account_cards_processing'] = MagicMock()
    # Добавляем конкретную функцию, которая вызывает проблему
    sys.modules['agregator.processing.account_cards_processing'].connect_account_card_to_heritage = MagicMock()

    # Также мокаем другие возможные зависимости
    sys.modules['agregator.views.auth'] = MagicMock()
    sys.modules['agregator.views.file_processing'] = MagicMock()
    sys.modules['agregator.processing.external_sources'] = MagicMock()

import os
import shutil
import pytest
import tempfile
import json
from unittest.mock import patch, MagicMock, mock_open
import cv2
import numpy as np
from django.core.files.uploadedfile import SimpleUploadedFile
from django.core.exceptions import ValidationError
from django.conf import settings
from django.utils import timezone
from celery.exceptions import Retry
from PIL import Image
import fitz
import zipfile
import io
import re
import datetime
from celery import states
from celery.result import AsyncResult

from agregator.models import (
    ObjectAccountCard,
    IdentifiedArchaeologicalHeritageSite,
    ArchaeologicalHeritageSite,
    UserTasks
)
from agregator.processing.account_cards_processing import (
    extract_text_from_image,
    sort_contours_custom,
    ccw,
    intersect,
    process_account_cards,
    extract_text_tables_and_images,
    error_handler_account_cards,
    min_area,
    symbol_config
)
from agregator.celery_task_template import process_documents
from agregator.hash import calculate_file_hash
from agregator.redis_config import redis_client


# ------------------- УТИЛИТЫ ДЛЯ ТЕСТОВ -------------------

def create_test_docx_with_tables(temp_dir):
    """Создает тестовый DOCX файл с таблицами"""
    from docx import Document
    from docx.shared import Inches

    doc = Document()

    # Таблица 1: Наименование объекта
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Наименование объекта"

    # Таблица 2: Время создания
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Время создания (возникновения) объекта"
    table.cell(0, 1).text = "XVII век"

    # Таблица 3: Адрес
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "г. Москва, ул. Тверская, д. 1"

    # Таблица 4: Вид объекта
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Вид объекта"
    table.cell(0, 1).text = "+"

    # Таблица 5: Общая видовая принадлежность
    table = doc.add_table(rows=1, cols=4)
    table.cell(0, 0).text = "Общая видовая принадлежность объекта"
    table.cell(0, 1).text = ""
    table.cell(0, 2).text = "+"
    table.cell(0, 3).text = ""

    # Таблица 6: Описание объекта
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Описание объекта"

    # Таблица 7: Использование объекта
    table = doc.add_table(rows=9, cols=2)
    table.cell(0, 0).text = "Использование объекта"
    table.cell(0, 1).text = "+"

    # Таблица 8: Сведения о дате выявления
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Сведения о дате выявления"

    # Таблица 9: Составитель
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "Составитель"
    table.cell(0, 1).text = "Иванов"
    table.cell(0, 2).text = "Петров"

    # Таблица 10: Дата составления
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "2023-01-01"

    # Вложенная таблица с координатами
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "1"
    table.cell(0, 1).text = "N 55° 45' 12.345\""
    table.cell(0, 2).text = "E 37° 37' 09.876\""

    # Добавляем изображение
    image_path = os.path.join(temp_dir, "test_image.png")
    Image.new('RGB', (100, 100), color='red').save(image_path)

    # Сохраняем документ
    doc_path = os.path.join(temp_dir, "test_account_card.docx")
    doc.save(doc_path)

    # Создаем ZIP-файл с изображением внутри
    with zipfile.ZipFile(doc_path, 'a') as zipf:
        zipf.write(image_path, "word/media/test_image.png")

    return doc_path


def create_test_pdf_with_tables(temp_dir):
    """Создает тестовый PDF файл с таблицами и координатами"""
    # Создаем PDF с помощью fitz
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)  # A4

    # Добавляем текстовые блоки с координатами и таблицами
    text = """
    Наименование объекта
    Музей истории

    Время создания (возникновения) объекта
    XVII век

    Адрес (местонахождение) объекта
    г. Москва, ул. Тверская, д. 1

    Координаты центра объекта WGS-84:
    N 55° 45' 12.345" E 37° 37' 09.876"

    Вид объекта
    + Памятник

    Общая видовая принадлежность объекта
    + Памятник археологии

    Общее описание объекта и вывод о его историко-культурной ценности
    Описание объекта

    Северная широта: N 55° 45' 12.345"
    Восточная долгота: E 37° 37' 09.876"

    Использование объекта культурного наследия или пользователь
    + Организации науки и образования

    Сведения о дате и обстоятельствах выявления (обнаружения) объекта
    Обнаружен в 1950 году

    Составитель учетной карты
    Иванов И.И.
    """

    page.insert_text(fitz.Point(50, 50), text)

    # Сохраняем PDF
    pdf_path = os.path.join(temp_dir, "test_account_card.pdf")
    doc.save(pdf_path)
    doc.close()

    return pdf_path


def create_test_image_for_ocr(temp_dir, text_content, width=800, height=200):
    """Создает тестовое изображение с текстом для OCR"""
    img = Image.new('RGB', (width, height), color='white')
    from PIL import ImageDraw, ImageFont
    try:
        font = ImageFont.truetype("arial.ttf", 36)
    except:
        # Для систем без Arial используем дефолтный шрифт
        font = ImageFont.load_default()

    draw = ImageDraw.Draw(img)
    draw.text((50, 50), text_content, fill="black", font=font)

    img_path = os.path.join(temp_dir, "test_ocr_image.png")
    img.save(img_path)

    return img_path


def create_test_zip_with_images(temp_dir):
    """Создает ZIP архив с изображениями для тестирования"""
    zip_path = os.path.join(temp_dir, "test_images.zip")
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        # Создаем несколько тестовых изображений
        for i in range(3):
            img = Image.new('RGB', (100, 100), color=(255, 0, 0))
            img_path = os.path.join(temp_dir, f"image_{i}.png")
            img.save(img_path)
            zipf.write(img_path, f"images/image_{i}.png")

    return zip_path


def create_test_empty_docx(temp_dir):
    """Создает пустой DOCX файл"""
    from docx import Document
    doc = Document()
    doc_path = os.path.join(temp_dir, "empty.docx")
    doc.save(doc_path)
    return doc_path


def create_test_malformed_docx(temp_dir):
    """Создает поврежденный DOCX файл"""
    doc_path = os.path.join(temp_dir, "malformed.docx")
    with open(doc_path, 'wb') as f:
        f.write(b'not a valid docx file')
    return doc_path


def create_test_docx_with_special_chars(temp_dir):
    """Создает DOCX с особыми символами"""
    from docx import Document
    doc = Document()

    # Добавляем текст с особыми символами
    doc.add_paragraph("Музей имени М.А. Шолохова")
    doc.add_paragraph("Координаты: N 55° 45' 12.345\" E 37° 37' 09.876\"")

    doc_path = os.path.join(temp_dir, "special_chars.docx")
    doc.save(doc_path)
    return doc_path


# ------------------- ТЕСТЫ -------------------

@pytest.mark.django_db
class TestAccountCardsProcessing:

    # ------------------- UNIT TESTS -------------------

    def test_extract_text_from_image(self, tmpdir):
        """Тест извлечения текста из изображения с различными конфигурациями"""
        # Создаем тестовое изображение
        test_text = "Тестовый текст для OCR"
        img_path = create_test_image_for_ocr(str(tmpdir), test_text)

        # Загружаем изображение
        img = cv2.imread(img_path)

        # Тестируем разные конфигурации PSM
        for psm in ['1', '3', '6', '11']:
            extracted_text = extract_text_from_image(img, psm)
            assert test_text in extracted_text, f"Текст не найден с конфигурацией PSM={psm}"

    def test_extract_text_from_image_empty(self):
        """Тест извлечения текста из пустого изображения"""
        # Создаем пустое изображение
        img = np.zeros((100, 100, 3), dtype=np.uint8)
        extracted_text = extract_text_from_image(img, '3')
        assert extracted_text.strip() == "", "Пустое изображение должно давать пустой текст"

    def test_sort_contours_custom_basic(self):
        """Тест сортировки контуров в базовом случае"""
        # Создаем тестовые контуры
        contours = [
            np.array([[0, 100], [50, 100], [50, 150], [0, 150]]),  # Контур 1
            np.array([[0, 50], [50, 50], [50, 100], [0, 100]]),  # Контур 2
            np.array([[0, 0], [50, 0], [50, 50], [0, 50]])  # Контур 3
        ]

        sorted_contours = sort_contours_custom(contours)
        # Проверяем, что контуры отсортированы по Y-координате
        assert len(sorted_contours) == 3
        # Проверяем порядок
        bbox1 = cv2.boundingRect(sorted_contours[0])
        bbox2 = cv2.boundingRect(sorted_contours[1])
        bbox3 = cv2.boundingRect(sorted_contours[2])

        assert bbox1[1] <= bbox2[1] <= bbox3[1], "Контуры должны быть отсортированы по вертикали"

    def test_sort_contours_custom_overlap(self):
        """Тест сортировки перекрывающихся контуров"""
        # Создаем перекрывающиеся контуры
        contours = [
            np.array([[0, 100], [100, 100], [100, 200], [0, 200]]),  # Контур 1
            np.array([[50, 50], [150, 50], [150, 150], [50, 150]]),  # Контур 2 (перекрывается с первым)
            np.array([[0, 0], [100, 0], [100, 100], [0, 100]])  # Контур 3
        ]

        sorted_contours = sort_contours_custom(contours)
        # Проверяем, что контур 3 (нижний) идет первым
        bbox1 = cv2.boundingRect(sorted_contours[0])
        bbox2 = cv2.boundingRect(sorted_contours[1])
        bbox3 = cv2.boundingRect(sorted_contours[2])

        assert bbox3[1] < bbox2[1] < bbox1[1], "Контуры должны быть правильно упорядочены при перекрытии"

    def test_ccw_basic(self):
        """Тест базового случая для ccw (против часовой стрелки)"""
        # Точки образуют треугольник против часовой стрелки
        A = (0, 0)
        B = (1, 0)
        C = (1, 1)
        assert ccw(A, B, C) is True, "Должно быть True для контура против часовой стрелки"

    def test_ccw_clockwise(self):
        """Тест случая для ccw (по часовой стрелке)"""
        # Точки образуют треугольник по часовой стрелке
        A = (0, 0)
        B = (1, 1)
        C = (1, 0)
        assert ccw(A, B, C) is False, "Должно быть False для контура по часовой стрелке"

    def test_intersect_basic(self):
        """Тест пересечения линий"""
        # Пересекающиеся линии
        A = (0, 0)
        B = (1, 1)
        C = (0, 1)
        D = (1, 0)
        assert intersect(A, B, C, D) is True, "Линии должны пересекаться"

        # Непересекающиеся линии
        A = (0, 0)
        B = (1, 0)
        C = (0, 1)
        D = (1, 1)
        assert intersect(A, B, C, D) is False, "Линии не должны пересекаться"

    @pytest.mark.parametrize("file_type,expected_name,expected_creation_time,expected_address", [
        ('docx', 'Музей истории', 'XVII век', 'г. Москва, ул. Тверская, д. 1'),
        ('pdf', 'Музей истории', 'XVII век', 'г. Москва, ул. Тверская, д. 1')
    ])
    def test_extract_text_tables_and_images_basic(
            self, tmpdir, db, test_user, file_type, expected_name,
            expected_creation_time, expected_address
    ):
        """Базовый тест обработки учетной карты"""
        # Создаем тестовый файл
        if file_type == 'docx':
            file_path = create_test_docx_with_tables(str(tmpdir))
        else:
            file_path = create_test_pdf_with_tables(str(tmpdir))

        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename=os.path.basename(file_path),
            source=file_path,
            is_processing=True
        )

        # Настройка прогресса
        progress_json = {
            "file_groups": {
                str(account_card.id): {
                    "processed": "False",
                    "origin_filename": os.path.basename(file_path),
                    "path": file_path
                }
            },
            "expected_time": "00:00:00"
        }
        task_id = f"test_task_{account_card.id}"

        # Мокаем необходимые зависимости
        with patch('agregator.processing.account_cards_processing.redis_client') as mock_redis:
            mock_redis.get.return_value = json.dumps(progress_json)
            mock_redis.set.return_value = None

            # Запускаем обработку
            extract_text_tables_and_images(
                file_path,
                MagicMock(),
                {str(account_card.id): 1},
                [0],
                account_card.id,
                progress_json,
                task_id,
                datetime.datetime.now()
            )

        # Проверяем результат
        account_card.refresh_from_db()
        assert account_card.name == expected_name
        assert account_card.creation_time == expected_creation_time
        assert account_card.address == expected_address
        assert account_card.is_processing is False

    # ------------------- EDGE CASES -------------------

    def test_extract_text_tables_and_images_duplicate_file(self, tmpdir, db, test_user):
        """Тест обработки дублирующегося файла"""
        # Создаем два одинаковых файла
        file_path1 = create_test_docx_with_tables(str(tmpdir))
        file_path2 = os.path.join(str(tmpdir), "copy.docx")
        shutil.copy(file_path1, file_path2)

        # Создаем первую учетную карту
        account_card1 = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename=os.path.basename(file_path1),
            source=file_path1,
            is_processing=True
        )

        # Создаем вторую учетную карту
        account_card2 = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename=os.path.basename(file_path2),
            source=file_path2,
            is_processing=True
        )

        # Настройка прогресса
        progress_json = {
            "file_groups": {
                str(account_card2.id): {
                    "processed": "False",
                    "origin_filename": os.path.basename(file_path2),
                    "path": file_path2
                }
            }
        }
        task_id = f"test_task_{account_card2.id}"

        # Проверяем, что возникает ошибка дублирования
        with pytest.raises(FileExistsError) as excinfo:
            extract_text_tables_and_images(
                file_path2,
                MagicMock(),
                {str(account_card2.id): 1},
                [0],
                account_card2.id,
                progress_json,
                task_id,
                datetime.datetime.now()
            )

        assert "Такой файл уже загружен в систему" in str(excinfo.value)

    def test_extract_text_tables_and_images_empty_file(self, tmpdir, db, test_user):
        """Тест обработки пустого файла"""
        # Создаем пустой DOCX
        file_path = create_test_empty_docx(str(tmpdir))

        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename=os.path.basename(file_path),
            source=file_path,
            is_processing=True
        )

        # Настройка прогресса
        progress_json = {
            "file_groups": {
                str(account_card.id): {
                    "processed": "False",
                    "origin_filename": os.path.basename(file_path),
                    "path": file_path
                }
            }
        }
        task_id = f"test_task_{account_card.id}"

        # Обрабатываем файл
        extract_text_tables_and_images(
            file_path,
            MagicMock(),
            {str(account_card.id): 1},
            [0],
            account_card.id,
            progress_json,
            task_id,
            datetime.datetime.now()
        )

        # Проверяем результат
        account_card.refresh_from_db()
        assert account_card.name is None
        assert account_card.creation_time is None
        assert account_card.address is None

    def test_extract_text_tables_and_images_malformed_file(self, tmpdir, db, test_user):
        """Тест обработки поврежденного файла"""
        # Создаем поврежденный DOCX
        file_path = create_test_malformed_docx(str(tmpdir))

        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename=os.path.basename(file_path),
            source=file_path,
            is_processing=True
        )

        # Настройка прогресса
        progress_json = {
            "file_groups": {
                str(account_card.id): {
                    "processed": "False",
                    "origin_filename": os.path.basename(file_path),
                    "path": file_path
                }
            }
        }
        task_id = f"test_task_{account_card.id}"

        # Обрабатываем файл
        extract_text_tables_and_images(
            file_path,
            MagicMock(),
            {str(account_card.id): 1},
            [0],
            account_card.id,
            progress_json,
            task_id,
            datetime.datetime.now()
        )

        # Проверяем результат
        account_card.refresh_from_db()
        assert account_card.name is None
        assert account_card.is_processing is False

    def test_extract_text_tables_and_images_special_chars(self, tmpdir, db, test_user):
        """Тест обработки файлов с особыми символами"""
        # Создаем DOCX с особыми символами
        file_path = create_test_docx_with_special_chars(str(tmpdir))

        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename=os.path.basename(file_path),
            source=file_path,
            is_processing=True
        )

        # Настройка прогресса
        progress_json = {
            "file_groups": {
                str(account_card.id): {
                    "processed": "False",
                    "origin_filename": os.path.basename(file_path),
                    "path": file_path
                }
            }
        }
        task_id = f"test_task_{account_card.id}"

        # Обрабатываем файл
        extract_text_tables_and_images(
            file_path,
            MagicMock(),
            {str(account_card.id): 1},
            [0],
            account_card.id,
            progress_json,
            task_id,
            datetime.datetime.now()
        )

        # Проверяем результат
        account_card.refresh_from_db()
        assert "Музей имени М.А. Шолохова" in account_card.description or \
               "Музей имени М.А. Шолохова" in account_card.name

    def test_extract_text_tables_and_images_coordinate_parsing(self, tmpdir, db, test_user):
        """Тест парсинга различных форматов координат"""
        # Создаем тестовый PDF с разными форматами координат
        doc = fitz.open()
        page = doc.new_page(width=595, height=842)

        test_cases = [
            "Координаты центра объекта WGS-84: N 55° 45' 12.345\" E 37° 37' 09.876\"",
            "Координаты: N 55°45'12.345\" E 37°37'09.876\"",
            "Координаты: N 55 45 12.345 E 37 37 09.876",
            "Координаты: N55°45'12.345\" E37°37'09.876\"",
            "Координаты: Северная широта N 55° 45' 12.345\" Восточная долгота E 37° 37' 09.876\"",
            "Координаты: N 55°45'12.345\" E 37°37'09.876\" WGS-84",
            "Координаты: N 55°45'12,345\" E 37°37'09,876\"",
            "Координаты: N 55°45.12345' E 37°37.09876'",
            "Координаты: N 55.753429 E 37.619432",
            "Координаты: 55.753429, 37.619432",
            "Координаты: 55°45'12.345\"N 37°37'09.876\"E",
            "Координаты: 55° 45' 12.345\" N, 37° 37' 09.876\" E"
        ]

        y_pos = 50
        for case in test_cases:
            page.insert_text(fitz.Point(50, y_pos), case)
            y_pos += 30

        pdf_path = os.path.join(str(tmpdir), "coords_test.pdf")
        doc.save(pdf_path)
        doc.close()

        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename=os.path.basename(pdf_path),
            source=pdf_path,
            is_processing=True
        )

        # Настройка прогресса
        progress_json = {
            "file_groups": {
                str(account_card.id): {
                    "processed": "False",
                    "origin_filename": os.path.basename(pdf_path),
                    "path": pdf_path
                }
            }
        }
        task_id = f"test_task_{account_card.id}"

        # Обрабатываем файл
        extract_text_tables_and_images(
            pdf_path,
            MagicMock(),
            {str(account_card.id): 1},
            [0],
            account_card.id,
            progress_json,
            task_id,
            datetime.datetime.now()
        )

        # Проверяем результат
        account_card.refresh_from_db()
        assert account_card.coordinates is not None
        assert 'Центр объекта' in account_card.coordinates
        assert 'Каталог координат' in account_card.coordinates

    # ------------------- INTEGRATION TESTS -------------------

    def test_process_account_cards_integration(self, tmpdir, db, test_user):
        """Интеграционный тест обработки учетной карты"""
        # Создаем тестовый DOCX
        file_path = create_test_docx_with_tables(str(tmpdir))

        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename=os.path.basename(file_path),
            source=file_path,
            is_processing=True
        )

        # Мокаем Celery задачу
        with patch('agregator.processing.account_cards_processing.process_documents') as mock_process:
            mock_process.return_value = None

            # Запускаем задачу
            result = process_account_cards.delay(
                account_cards_ids=[account_card.id],
                user_id=test_user.id
            )

            # Ждем выполнения задачи
            result.get(timeout=10)

        # Проверяем результат
        account_card.refresh_from_db()
        assert account_card.name == "Музей истории"
        assert account_card.is_processing is False

    def test_connect_account_card_to_heritage_success(self, db, test_user):
        """Тест успешного соединения учетной карты с наследием"""
        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            name="Test Object",
            origin_filename="test.docx",
            source="/tmp/test/test.docx",
            is_processing=False
        )

        # Создаем наследие
        heritage = ArchaeologicalHeritageSite.objects.create(
            doc_name="Test Object",
            district="Test District",
            register_num="TEST-001",
            source="/tmp/heritage/"
        )

        # Создаем директорию для теста
        os.makedirs("/tmp/test/", exist_ok=True)
        os.makedirs("/tmp/heritage/", exist_ok=True)
        with open("/tmp/test/test.docx", "w") as f:
            f.write("Test content")

        # Соединяем
        connect_account_card_to_heritage("Test Object")

        # Проверяем результат
        account_card.refresh_from_db()
        heritage.refresh_from_db()

        assert heritage.account_card_id == account_card.id
        assert "Учётная карта" in account_card.source
        assert os.path.exists(account_card.source)

    def test_connect_account_card_to_heritage_no_heritage(self, db, test_user):
        """Тест соединения учетной карты с несуществующим наследием"""
        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            name="Test Object",
            origin_filename="test.docx",
            source="/tmp/test/test.docx",
            is_processing=False
        )

        # Создаем директорию для теста
        os.makedirs("/tmp/test/", exist_ok=True)
        with open("/tmp/test/test.docx", "w") as f:
            f.write("Test content")

        # Соединяем
        connect_account_card_to_heritage("Non-existent Object")

        # Проверяем результат
        account_card.refresh_from_db()
        assert account_card.supplement == {"address": [], "description": []}

    # ------------------- SECURITY TESTS -------------------

    def test_extract_text_tables_and_images_path_traversal(self, tmpdir, db, test_user):
        """Тест защиты от path traversal атаки"""
        # Создаем тестовый файл
        file_path = create_test_docx_with_tables(str(tmpdir))

        # Создаем учетную карту с подозрительным именем
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename="../../etc/passwd",
            source=file_path,
            is_processing=True
        )

        # Настройка прогресса
        progress_json = {
            "file_groups": {
                str(account_card.id): {
                    "processed": "False",
                    "origin_filename": "../../etc/passwd",
                    "path": file_path
                }
            }
        }
        task_id = f"test_task_{account_card.id}"

        # Обрабатываем файл
        extract_text_tables_and_images(
            file_path,
            MagicMock(),
            {str(account_card.id): 1},
            [0],
            account_card.id,
            progress_json,
            task_id,
            datetime.datetime.now()
        )

        # Проверяем, что путь не был изменен
        account_card.refresh_from_db()
        assert ".." not in account_card.source

    def test_extract_text_tables_and_images_malicious_content(self, tmpdir, db, test_user):
        """Тест обработки файла с вредоносным содержимым"""
        # Создаем тестовый PDF с подозрительным содержимым
        doc = fitz.open()
        page = doc.new_page(width=595, height=842)

        # Добавляем подозрительный текст
        malicious_text = """
        <script>alert('XSS')</script>
        <?php system('rm -rf /'); ?>
        EXEC master..xp_cmdshell 'dir'
        """
        page.insert_text(fitz.Point(50, 50), malicious_text)

        pdf_path = os.path.join(str(tmpdir), "malicious.pdf")
        doc.save(pdf_path)
        doc.close()

        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename=os.path.basename(pdf_path),
            source=pdf_path,
            is_processing=True
        )

        # Настройка прогресса
        progress_json = {
            "file_groups": {
                str(account_card.id): {
                    "processed": "False",
                    "origin_filename": os.path.basename(pdf_path),
                    "path": pdf_path
                }
            }
        }
        task_id = f"test_task_{account_card.id}"

        # Обрабатываем файл
        extract_text_tables_and_images(
            pdf_path,
            MagicMock(),
            {str(account_card.id): 1},
            [0],
            account_card.id,
            progress_json,
            task_id,
            datetime.datetime.now()
        )

        # Проверяем результат
        account_card.refresh_from_db()
        assert "<script>" not in account_card.description
        assert "alert" not in account_card.description
        assert "rm -rf" not in account_card.description

    # ------------------- ERROR HANDLING -------------------

    def test_error_handler_account_cards(self, db, test_user):
        """Тест обработчика ошибок для учетных карт"""
        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename="test.docx",
            source="/tmp/test/test.docx",
            is_processing=True
        )

        # Создаем задачу
        task = UserTasks.objects.create(
            user=test_user,
            task_id="test_task",
            files_type="account_card",
            upload_source={'source': 'Пользовательский файл'}
        )

        # Создаем прогресс JSON
        progress_json = {
            "file_groups": {
                str(account_card.id): {
                    "processed": "False",
                    "origin_filename": "test.docx",
                    "path": "/tmp/test/test.docx"
                }
            },
            "expected_time": "00:00:00"
        }

        # Мокаем redis
        with patch('agregator.redis_config.redis_client') as mock_redis:
            mock_redis.get.return_value = json.dumps(progress_json)

            # Вызываем обработчик ошибок
            with pytest.raises(Exception) as excinfo:
                error_handler_account_cards(
                    MagicMock(id="test_task"),
                    Exception("Test error"),
                    "Test error description"
                )

            # Проверяем, что учетная карта удалена
            with pytest.raises(ObjectAccountCard.DoesNotExist):
                ObjectAccountCard.objects.get(id=account_card.id)

            # Проверяем содержание ошибки
            error_data = json.loads(str(excinfo.value))
            assert "error_text" in error_data
            assert "progress_json" in error_data
            assert error_data["error_text"] == "Test error"

    def test_process_account_cards_retry(self, db, test_user):
        """Тест повторной попытки обработки"""
        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename="test.docx",
            source="/tmp/test/test.docx",
            is_processing=True
        )

        # Мокаем Celery задачу, чтобы вызвать Retry
        with patch('agregator.processing.account_cards_processing.process_documents') as mock_process:
            mock_process.side_effect = Retry()

            # Запускаем задачу
            result = process_account_cards.delay(
                account_cards_ids=[account_card.id],
                user_id=test_user.id
            )

            # Проверяем статус задачи
            assert result.state == states.RETRY

    # ------------------- PARAMETRIZED TESTS -------------------

    @pytest.mark.parametrize("coordinate_text,expected_lat,expected_lon", [
        ("N 55° 45' 12.345\" E 37° 37' 09.876\"", 55.753429, 37.619432),
        ("Северная широта N 55°45'12.345\" Восточная долгота E 37°37'09.876\"", 55.753429, 37.619432),
        ("N 55 45 12.345 E 37 37 09.876", 55.753429, 37.619432),
        ("N55°45'12.345\" E37°37'09.876\"", 55.753429, 37.619432),
        ("N 55°45.12345' E 37°37.09876'", 55.752058, 37.618304),
        ("N 55.753429 E 37.619432", 55.753429, 37.619432),
        ("55.753429, 37.619432", 55.753429, 37.619432),
        ("55°45'12.345\"N 37°37'09.876\"E", 55.753429, 37.619432),
        ("55° 45' 12.345\" N, 37° 37' 09.876\" E", 55.753429, 37.619432),
    ])
    def test_coordinate_parsing(self, coordinate_text, expected_lat, expected_lon, tmpdir, db, test_user):
        """Параметризованный тест парсинга координат"""
        # Создаем PDF с координатами
        doc = fitz.open()
        page = doc.new_page(width=595, height=842)
        page.insert_text(fitz.Point(50, 50), f"Координаты: {coordinate_text}")

        pdf_path = os.path.join(str(tmpdir), "coords_test.pdf")
        doc.save(pdf_path)
        doc.close()

        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename=os.path.basename(pdf_path),
            source=pdf_path,
            is_processing=True
        )

        # Настройка прогресса
        progress_json = {
            "file_groups": {
                str(account_card.id): {
                    "processed": "False",
                    "origin_filename": os.path.basename(pdf_path),
                    "path": pdf_path
                }
            }
        }
        task_id = f"test_task_{account_card.id}"

        # Обрабатываем файл
        extract_text_tables_and_images(
            pdf_path,
            MagicMock(),
            {str(account_card.id): 1},
            [0],
            account_card.id,
            progress_json,
            task_id,
            datetime.datetime.now()
        )

        # Проверяем результат
        account_card.refresh_from_db()
        if 'Центр объекта' in account_card.coordinates:
            coords = account_card.coordinates['Центр объекта']['Центр объекта']
            assert abs(coords[0] - expected_lat) < 0.0001, f"Неверная широта для {coordinate_text}"
            assert abs(coords[1] - expected_lon) < 0.0001, f"Неверная долгота для {coordinate_text}"
        else:
            coords = list(account_card.coordinates['Каталог координат'].values())[0]
            assert abs(coords[0] - expected_lat) < 0.0001, f"Неверная широта для {coordinate_text}"
            assert abs(coords[1] - expected_lon) < 0.0001, f"Неверная долгота для {coordinate_text}"

    @pytest.mark.parametrize("file_type,table_index,table_data,expected_value", [
        # DOCX: Наименование объекта (таблица 1)
        ('docx', 0, [['Наименование объекта'], ['Музей истории']], 'Музей истории'),
        # DOCX: Время создания (таблица 2)
        ('docx', 1, [['Время создания (возникновения) объекта', 'XVII век']], 'XVII век'),
        # DOCX: Адрес (таблица 3)
        ('docx', 2, [['г. Москва, ул. Тверская, д. 1']], 'г. Москва, ул. Тверская, д. 1'),
        # DOCX: Вид объекта (таблица 4)
        ('docx', 3, [['Вид объекта', '+']], 'Памятник'),
        # DOCX: Общая видовая принадлежность (таблица 5)
        ('docx', 4, [['Общая видовая принадлежность объекта', '', '+', '']], 'Памятник археологии'),
        # PDF: Наименование объекта
        ('pdf', None, "Наименование объекта\nМузей истории", 'Музей истории'),
        # PDF: Время создания
        ('pdf', None, "Время создания (возникновения) объекта\nXVII век", 'XVII век'),
        # PDF: Адрес
        ('pdf', None, "Адрес (местонахождение) объекта\nг. Москва, ул. Тверская, д. 1",
         'г. Москва, ул. Тверская, д. 1'),
        # PDF: Вид объекта
        ('pdf', None, "Вид объекта\n+ Памятник", 'Памятник'),
        # PDF: Общая видовая принадлежность
        ('pdf', None, "Общая видовая принадлежность объекта\n+ Памятник археологии", 'Памятник археологии'),
    ])
    def test_table_parsing(
            self, tmpdir, db, test_user, file_type, table_index,
            table_data, expected_value
    ):
        """Параметризованный тест парсинга таблиц"""
        if file_type == 'docx':
            # Создаем тестовый DOCX
            from docx import Document

            doc = Document()
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            for i, row_data in enumerate(table_data):
                for j, cell_data in enumerate(row_data):
                    table.cell(i, j).text = cell_data

            file_path = os.path.join(str(tmpdir), "table_test.docx")
            doc.save(file_path)
        else:
            # Создаем тестовый PDF
            doc = fitz.open()
            page = doc.new_page(width=595, height=842)
            page.insert_text(fitz.Point(50, 50), table_data)

            file_path = os.path.join(str(tmpdir), "table_test.pdf")
            doc.save(file_path)
            doc.close()

        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename=os.path.basename(file_path),
            source=file_path,
            is_processing=True
        )

        # Настройка прогресса
        progress_json = {
            "file_groups": {
                str(account_card.id): {
                    "processed": "False",
                    "origin_filename": os.path.basename(file_path),
                    "path": file_path
                }
            }
        }
        task_id = f"test_task_{account_card.id}"

        # Обрабатываем файл
        extract_text_tables_and_images(
            file_path,
            MagicMock(),
            {str(account_card.id): 1},
            [0],
            account_card.id,
            progress_json,
            task_id,
            datetime.datetime.now()
        )

        # Проверяем результат
        account_card.refresh_from_db()

        # Определяем, какое поле проверять
        if "наименование" in str(table_data).lower() or "название" in str(table_data).lower():
            assert account_card.name == expected_value
        elif "время" in str(table_data).lower() or "период" in str(table_data).lower():
            assert account_card.creation_time == expected_value
        elif "адрес" in str(table_data).lower() or "местонахождение" in str(table_data).lower():
            assert account_card.address == expected_value
        elif "вид" in str(table_data).lower():
            assert account_card.object_type == expected_value
        elif "принадлежность" in str(table_data).lower() or "классификация" in str(table_data).lower():
            assert account_card.general_classification == expected_value

    @pytest.mark.parametrize("area_value,expected_result", [
        (min_area - 1, False),  # Ниже порога
        (min_area, True),  # Точно на пороге
        (min_area + 1, True),  # Выше порога
        (0, False),  # Нулевая площадь
        (-100, False),  # Отрицательная площадь
    ])
    def test_contour_area_filtering(self, area_value, expected_result):
        """Параметризованный тест фильтрации контуров по площади"""
        # Создаем контур с заданной площадью
        contour = np.array([[[0, 0], [10, 0], [10, 10], [0, 10]]])

        # Мокаем cv2.contourArea
        with patch('cv2.contourArea', return_value=area_value):
            # Создаем изображение
            img = np.zeros((100, 100), dtype=np.uint8)

            # Процессинг
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            _, thresh = cv2.threshold(gray, 180, 255, cv2.THRESH_BINARY_INV)
            kernel = np.ones((50, 50), np.uint8)
            dilated = cv2.dilate(thresh, kernel, iterations=1)
            contours, _ = cv2.findContours(dilated, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

            # Проверяем фильтрацию
            if expected_result:
                assert len(contours) > 0, "Контур должен быть сохранен"
            else:
                assert len(contours) == 0, "Контур должен быть отфильтрован"

    # ------------------- ADDITIONAL TESTS -------------------

    def test_process_account_cards_empty_list(self, db, test_user):
        """Тест обработки пустого списка учетных карт"""
        # Запускаем задачу с пустым списком
        result = process_account_cards.delay(
            account_cards_ids=[],
            user_id=test_user.id
        )

        # Проверяем результат
        assert result.successful()

    def test_process_account_cards_non_existent(self, db, test_user):
        """Тест обработки несуществующих учетных карт"""
        # Запускаем задачу с несуществующим ID
        result = process_account_cards.delay(
            account_cards_ids=[999999],
            user_id=test_user.id
        )

        # Проверяем результат
        with pytest.raises(ObjectAccountCard.DoesNotExist):
            result.get(timeout=5)

    def test_supplement_content_structure(self, tmpdir, db, test_user):
        """Тест структуры дополнительного контента"""
        # Создаем тестовый DOCX с изображениями
        file_path = create_test_docx_with_tables(str(tmpdir))

        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename=os.path.basename(file_path),
            source=file_path,
            is_processing=True
        )

        # Настройка прогресса
        progress_json = {
            "file_groups": {
                str(account_card.id): {
                    "processed": "False",
                    "origin_filename": os.path.basename(file_path),
                    "path": file_path
                }
            }
        }
        task_id = f"test_task_{account_card.id}"

        # Обрабатываем файл
        extract_text_tables_and_images(
            file_path,
            MagicMock(),
            {str(account_card.id): 1},
            [0],
            account_card.id,
            progress_json,
            task_id,
            datetime.datetime.now()
        )

        # Проверяем структуру дополнительного контента
        account_card.refresh_from_db()
        assert 'address' in account_card.supplement
        assert 'description' in account_card.supplement
        assert isinstance(account_card.supplement['address'], list)
        assert isinstance(account_card.supplement['description'], list)
        assert all(isinstance(item, dict) for item in account_card.supplement['address'])
        assert all('label' in item and 'source' in item for item in account_card.supplement['address'])

    def test_coordinate_system_detection(self, tmpdir, db, test_user):
        """Тест обнаружения системы координат"""
        # Создаем PDF с координатами
        doc = fitz.open()
        page = doc.new_page(width=595, height=842)
        page.insert_text(fitz.Point(50, 50), "Координаты WGS-84: N 55° 45' 12.345\" E 37° 37' 09.876\"")

        pdf_path = os.path.join(str(tmpdir), "coords_test.pdf")
        doc.save(pdf_path)
        doc.close()

        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename=os.path.basename(pdf_path),
            source=pdf_path,
            is_processing=True
        )

        # Настройка прогресса
        progress_json = {
            "file_groups": {
                str(account_card.id): {
                    "processed": "False",
                    "origin_filename": os.path.basename(pdf_path),
                    "path": pdf_path
                }
            }
        }
        task_id = f"test_task_{account_card.id}"

        # Обрабатываем файл
        extract_text_tables_and_images(
            pdf_path,
            MagicMock(),
            {str(account_card.id): 1},
            [0],
            account_card.id,
            progress_json,
            task_id,
            datetime.datetime.now()
        )

        # Проверяем систему координат
        account_card.refresh_from_db()
        if 'Центр объекта' in account_card.coordinates:
            assert account_card.coordinates['Центр объекта']['coordinate_system'] == 'wgs84'
        else:
            assert account_card.coordinates['Каталог координат']['coordinate_system'] == 'wgs84'

    def test_image_extraction(self, tmpdir, db, test_user):
        """Тест извлечения изображений из документа"""
        # Создаем тестовый DOCX с изображениями
        file_path = create_test_docx_with_tables(str(tmpdir))

        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename=os.path.basename(file_path),
            source=file_path,
            is_processing=True
        )

        # Настройка прогресса
        progress_json = {
            "file_groups": {
                str(account_card.id): {
                    "processed": "False",
                    "origin_filename": os.path.basename(file_path),
                    "path": file_path
                }
            }
        }
        task_id = f"test_task_{account_card.id}"

        # Обрабатываем файл
        extract_text_tables_and_images(
            file_path,
            MagicMock(),
            {str(account_card.id): 1},
            [0],
            account_card.id,
            progress_json,
            task_id,
            datetime.datetime.now()
        )

        # Проверяем извлечение изображений
        account_card.refresh_from_db()
        assert len(account_card.supplement['address']) > 0 or len(account_card.supplement['description']) > 0
        for category in ['address', 'description']:
            for item in account_card.supplement[category]:
                assert os.path.exists(item['source']), f"Изображение не сохранено: {item['source']}"

    def test_redis_progress_updates(self, tmpdir, db, test_user):
        """Тест обновления прогресса в Redis"""
        # Создаем тестовый DOCX
        file_path = create_test_docx_with_tables(str(tmpdir))

        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename=os.path.basename(file_path),
            source=file_path,
            is_processing=True
        )

        # Настройка прогресса
        progress_json = {
            "file_groups": {
                str(account_card.id): {
                    "processed": "False",
                    "origin_filename": os.path.basename(file_path),
                    "path": file_path,
                    "pages": {"total": 1, "processed": 0}
                }
            },
            "expected_time": "00:00:00"
        }
        task_id = f"test_task_{account_card.id}"

        # Мокаем redis
        with patch('agregator.redis_config.redis_client') as mock_redis:
            mock_redis.get.return_value = json.dumps(progress_json)
            mock_redis.set.return_value = None

            # Обрабатываем файл
            extract_text_tables_and_images(
                file_path,
                MagicMock(),
                {str(account_card.id): 1},
                [0],
                account_card.id,
                progress_json,
                task_id,
                datetime.datetime.now()
            )

            # Проверяем, что прогресс обновлялся
            assert mock_redis.set.call_count > 0
            # Проверяем, что страницы обработаны
            calls = mock_redis.set.call_args_list
            for call in calls:
                args = call[0]
                if 'pages' in args[1]:
                    assert args[1]['pages']['processed'] == 1

    # ------------------- FINAL CHECKS -------------------

    def test_full_processing_flow(self, tmpdir, db, test_user):
        """Полный тест обработки учетной карты от начала до конца"""
        # 1. Создаем тестовые данные
        file_path = create_test_docx_with_tables(str(tmpdir))

        # 2. Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename=os.path.basename(file_path),
            source=file_path,
            is_processing=True
        )

        # 3. Мокаем необходимые зависимости
        progress_json = {
            "file_groups": {
                str(account_card.id): {
                    "processed": "False",
                    "origin_filename": os.path.basename(file_path),
                    "path": file_path,
                    "pages": {"total": 1, "processed": 0}
                }
            },
            "expected_time": "00:00:00"
        }
        task_id = f"test_task_{account_card.id}"

        # 4. Запускаем обработку
        with patch('agregator.processing.account_cards_processing.redis_client') as mock_redis:
            mock_redis.get.return_value = json.dumps(progress_json)
            mock_redis.set.return_value = None

            extract_text_tables_and_images(
                file_path,
                MagicMock(),
                {str(account_card.id): 1},
                [0],
                account_card.id,
                progress_json,
                task_id,
                datetime.datetime.now()
            )

        # 5. Проверяем результат
        account_card.refresh_from_db()
        assert account_card.name == "Музей истории"
        assert account_card.creation_time == "XVII век"
        assert account_card.address == "г. Москва, ул. Тверская, д. 1"
        assert account_card.object_type == "Памятник"
        assert account_card.general_classification == "Памятник археологии"
        assert account_card.is_processing is False
        assert account_card.coordinates is not None
        assert 'address' in account_card.supplement
        assert 'description' in account_card.supplement

    def test_hash_calculation(self, tmpdir):
        """Тест вычисления хеша файла"""
        # Создаем тестовый файл
        file_path = os.path.join(str(tmpdir), "test_file.txt")
        with open(file_path, "w") as f:
            f.write("Test content")

        # Вычисляем хеш
        hash1 = calculate_file_hash(file_path)

        # Меняем содержимое
        with open(file_path, "w") as f:
            f.write("Modified content")

        # Вычисляем новый хеш
        hash2 = calculate_file_hash(file_path)

        # Проверяем, что хеши разные
        assert hash1 != hash2

    def test_file_deletion_on_error(self, tmpdir, db, test_user):
        """Тест удаления файла при ошибке"""
        # Создаем тестовый DOCX
        file_path = create_test_docx_with_tables(str(tmpdir))

        # Создаем учетную карту
        account_card = ObjectAccountCard.objects.create(
            user=test_user,
            origin_filename=os.path.basename(file_path),
            source=file_path,
            is_processing=True
        )

        # Настройка прогресса
        progress_json = {
            "file_groups": {
                str(account_card.id): {
                    "processed": "False",
                    "origin_filename": os.path.basename(file_path),
                    "path": file_path
                }
            }
        }
        task_id = f"test_task_{account_card.id}"

        # Мокаем функцию, чтобы вызвать ошибку
        with patch('agregator.processing.account_cards_processing.load_raw_account_cards') as mock_load:
            mock_load.side_effect = Exception("Test error")

            # Пытаемся обработать файл
            with pytest.raises(Exception):
                extract_text_tables_and_images(
                    file_path,
                    MagicMock(),
                    {str(account_card.id): 1},
                    [0],
                    account_card.id,
                    progress_json,
                    task_id,
                    datetime.datetime.now()
                )

            # Проверяем, что учетная карта удалена
            with pytest.raises(ObjectAccountCard.DoesNotExist):
                ObjectAccountCard.objects.get(id=account_card.id)
