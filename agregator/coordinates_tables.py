import re

import openpyxl
from docx import Document
import pdfplumber
import pandas as pd
from pyproj import Proj, transform

COORDINATE_SYSTEMS = [
    r'wgs.*?\d+',
    r'мск.*?\d+',
    r'гск.*?\d+',
]
COORDINATE_MARKS = {
    ('север', 'вост'): False,
    ('широт', 'долг'): False,
    ('x', 'y'): False,
    ('n', 'e'): False,
}
COORDINATE_TYPES = [
    r'[NS]+\d+°\s*\d+\'\s*\d+[\.,]\d+"|\d+°\s*\d+\'\s*\d+[\.,]\d+"[СЮ]+',
    r'[EW]+\d+°\s*\d+\'\s*\d+[\.,]\d+"|\d+°\s*\d+\'\s*\d+[\.,]\d+"[ВЗ]+',
    r'\d+°\s*\d+\'\s*\d+[\.,]\d+"',

    r'\d+[.,]+\d+',
]

projections = {
    "wgs84": Proj(proj='latlong', datum='WGS84'),
    "мск162": Proj(
        '+proj=tmerc +lat_0=0 +lon_0=78 +k=1 +x_0=62728.20 +y_0=-7546013.80 +ellps=krass +towgs84=23.57,-140.95,-79.8,0,0.35,0.79,-0.22 +units=m +no_defs'),
    "мск163": Proj(
        '+proj=tmerc +lat_0=0 +lon_0=81 +k=1 +x_0=65425.70 +y_0=-7434483.20 +ellps=krass +towgs84=23.57,-140.95,-79.8,0,0.35,0.79,-0.22 +units=m +no_defs'),
    "мск164": Proj(
        '+proj=tmerc +lat_0=0 +lon_0=84 +k=1 +x_0=86209.80 +y_0=-6542783.50 +ellps=krass +towgs84=23.57,-140.95,-79.8,0,0.35,0.79,-0.22 +units=m +no_defs'),
    "мск165": Proj(
        '+proj=tmerc +lat_0=0 +lon_0=87 +k=1 +x_0=105295.80 +y_0=-5652185.00 +ellps=krass +towgs84=23.57,-140.95,-79.8,0,0.35,0.79,-0.22 +units=m +no_defs'),
    "мск166": Proj(
        '+proj=tmerc +lat_0=0 +lon_0=90 +k=1 +x_0=107543.30 +y_0=-5540944.50 +ellps=krass +towgs84=23.57,-140.95,-79.8,0,0.35,0.79,-0.22 +units=m +no_defs'),
    "мск167": Proj(
        '+proj=tmerc +lat_0=0 +lon_0=93 +k=1 +x_0=106797.80 +y_0=-5578022.50 +ellps=krass +towgs84=23.57,-140.95,-79.8,0,0.35,0.79,-0.22 +units=m +no_defs'),
    "мск168": Proj(
        '+proj=tmerc +lat_0=0 +lon_0=96 +k=1 +x_0=108285.20 +y_0=-5503868.60 +ellps=krass +towgs84=23.57,-140.95,-79.8,0,0.35,0.79,-0.22 +units=m +no_defs'),
    "мск169": Proj(
        '+proj=tmerc +lat_0=0 +lon_0=99 +k=1 +x_0=117308.60 +y_0=-5503868.60 +ellps=krass +towgs84=23.57,-140.95,-79.8,0,0.35,0.79,-0.22 +units=m +no_defs'),
    "мск170": Proj(
        '+proj=tmerc +lat_0=0 +lon_0=102 +k=1 +x_0=90338.90 +y_0=-6357146.10 +ellps=krass +towgs84=23.57,-140.95,-79.8,0,0.35,0.79,-0.22 +units=m +no_defs'),
    "мск171": Proj(
        '+proj=tmerc +lat_0=0 +lon_0=105 +k=1 +x_0=87870.50 +y_0=-6468522.70 +ellps=krass +towgs84=23.57,-140.95,-79.8,0,0.35,0.79,-0.22 +units=m +no_defs'),
    "мск24зона1": Proj(
        '+proj=tmerc +lat_0=0 +lon_0=81.51666667 +k=1 +x_0=1500000 +y_0=-5416586.442 +ellps=krass +towgs84=23.57,-140.95,-79.8,0,0.35,0.79,-0.22 +units=m +no_defs'),
    "мск24зона2": Proj(
        '+proj=tmerc +lat_0=0 +lon_0=87.51666667 +k=1 +x_0=2500000 +y_0=-5416586.442 +ellps=krass +towgs84=23.57,-140.95,-79.8,0,0.35,0.79,-0.22 +units=m +no_defs'),
    "мск24зона3": Proj(
        '+proj=tmerc +lat_0=0 +lon_0=93.51666667 +k=1 +x_0=3500000 +y_0=-5416586.442 +ellps=krass +towgs84=23.57,-140.95,-79.8,0,0.35,0.79,-0.22 +units=m +no_defs'),
    "мск24зона4": Proj(
        '+proj=tmerc +lat_0=0 +lon_0=99.51666667 +k=1 +x_0=4500000 +y_0=-5416586.442 +ellps=krass +towgs84=23.57,-140.95,-79.8,0,0.35,0.79,-0.22 +units=m +no_defs'),
    "мск24зона5": Proj(
        '+proj=tmerc +lat_0=0 +lon_0=105.51666667 +k=1 +x_0=5500000 +y_0=-5416586.442 +ellps=krass +towgs84=23.57,-140.95,-79.8,0,0.35,0.79,-0.22 +units=m +no_defs'),
    "мск24зона6": Proj(
        '+proj=tmerc +lat_0=0 +lon_0=111.51666667 +k=1 +x_0=6500000 +y_0=-5416586.442 +ellps=krass +towgs84=23.57,-140.95,-79.8,0,0.35,0.79,-0.22 +units=m +no_defs'),
    "гск2011": Proj(
        '+proj=tmerc +lat_0=0 +lon_0=81 +k=1 +x_0=14500000 +y_0=0 +ellps=GSK2011 +towgs84=0,0,0,0,0,0,0 +units=m +no_defs +type=crs'),
}


def convert_to_wgs84(x, y, system):
    y, x = transform(projections[system], projections['wgs84'], y, x)
    return x, y


def convert_proj4(x, y, init_system, final_system):
    y, x = transform(projections[init_system], projections[final_system], y, x)
    return x, y


def str_is_float(string):
    try:
        result = float(string)
    except ValueError:
        return False
    return True


def normalize_coordinates(coord: str) -> str:
    coord = coord.strip()
    coord = coord.replace(' 0', ' ').replace(' ', '"')
    coord = coord[1:] if coord[0] == '0' else coord
    return coord


def dms_to_decimal(dms):
    """Преобразует координаты из формата DMS в десятичный формат."""
    dms = dms.replace(',', '.')
    dms = re.sub(r'[^\d°\'".]', '', dms)
    print('!!' + str(dms))
    parts = re.split('[°\'"]+', dms)
    print('!!' + str(parts))

    degrees = float(parts[0])
    minutes = float(parts[1])
    seconds = float(parts[2])

    decimal = degrees + minutes / 60 + seconds / 3600
    return decimal


def analyze_coordinates_in_tables_from_pdf(tables, file_path):
    """Анализирует таблицы на наличие координатных данных"""
    dfs = []
    full_text = extract_text_from_pdf(file_path)

    if not tables:
        return None, None, None

    last_num = last_width = last_target_cell = last_number_column = None
    appending = False

    coordinate_systems = []
    multiple_coord_sys = False
    for table_idx, table in enumerate(tables, 1):
        # Преобразуем таблицу в DataFrame
        df = pd.DataFrame(table)

        class CellMock:
            def __init__(self, value, row, col):
                self.value = value
                self.row = row + 1  # openpyxl использует 1-индексацию
                self.column = col + 1
                self.column_letter = openpyxl.utils.get_column_letter(col + 1)

        found_longitude = found_latitude = False
        legend_length = 0
        target_cell = number_column = None
        current_area = [None]

        # Анализируем все ячейки таблицы
        for row_idx in range(len(df)):
            for col_idx in range(len(df.columns)):
                cell_value = df.iat[row_idx, col_idx]
                if pd.isna(cell_value) or cell_value == "":
                    continue

                cell_str = str(cell_value).lower().strip()
                cell_mock = CellMock(cell_value, row_idx, col_idx)

                # Поиск систем координат
                for pattern in COORDINATE_SYSTEMS:
                    coord_sys = re.search(pattern, cell_str, re.IGNORECASE)
                    if coord_sys:
                        coord_sys = coord_sys.group(0)
                        if coord_sys not in coordinate_systems:
                            coordinate_systems.append(coord_sys)

                # Поиск меток координат
                for key in COORDINATE_MARKS.keys():
                    for pattern in key:
                        if len(cell_str) > 1 and re.search(pattern, cell_str, re.IGNORECASE) or cell_str == pattern:
                            COORDINATE_MARKS[key] = True

                # Поиск заголовков координат
                if cell_str in ('северная', 'широта', 'x', 'n') or 'север' in cell_str or 'широт' in cell_str:
                    found_latitude = True
                    legend_length += 1
                    target_cell = cell_mock
                elif cell_str in ('восточная', 'долгота', 'y', 'e') or 'вост' in cell_str or 'долг' in cell_str:
                    found_longitude = True
                    legend_length += 1
                    target_cell = cell_mock
                elif 'обознач' in cell_str or 'номер' in cell_str or '№' in cell_str:
                    legend_length += 1
                    number_column = col_idx

        for pattern in COORDINATE_SYSTEMS:
            coord_sys = re.search(pattern, full_text, re.IGNORECASE)
            if coord_sys:
                coord_sys = coord_sys.group(0).lower()
                if coord_sys not in coordinate_systems:
                    coordinate_systems.append(coord_sys)

        # Проверка на множественные системы координат
        if sum([1 for x in COORDINATE_MARKS.values() if x is True]) > 1 and len(coordinate_systems) > 1:
            multiple_coord_sys = True

        print('multiple_coord_sys: ' + str(multiple_coord_sys))
        print('coordinate_systems: ' + str(coordinate_systems))
        # Если нашли координаты, извлекаем подтаблицу
        print(str(found_latitude) + ' ' + str(found_longitude) + ' ' + str(target_cell))
        print(str(last_width) + ' ' + str(len(df.columns)))
        print('isinstance(last_num, str) and last_num.isdigit(): ' + str(
            isinstance(last_num, str) and last_num.isdigit()))
        print(type(last_num))
        print(df)
        print(appending)

        length, width = df.shape

        if last_number_column and last_number_column < len(df.columns) and isinstance(df.iloc[0, last_number_column],
                                                                                      str):
            print('str.isdigit(df.iloc[0, last_number_column]: ' + str(str.isdigit(df.iloc[0, last_number_column])))
            print(df)
            print(df.iloc[0, last_number_column])
        print('last_number_column: ' + str(last_number_column))

        if width is not None and last_number_column is not None and width >= last_number_column and length > 1:
            pass
            # print('Stykovochnye nomera: ' + str(last_num) + ' ' + str(df.iloc[0, last_number_column]) + ' ' + str(
            #    df.iloc[1, last_number_column]))
        print('last_width:' + str(last_width))
        print('len(df.columns):' + str(len(df.columns)))
        print('isinstance(last_num, str):' + str(isinstance(last_num, str)))
        if isinstance(last_num, str):
            print('last_num.isdigit():' + str(last_num.isdigit()))
        print('last_number_column:' + str(last_number_column))
        print('length:' + str(length))
        print('width:' + str(width))
        print('last_number_column:' + str(last_number_column))
        if last_number_column is not None and last_number_column < len(df.columns):
            print('df.iloc[0, last_number_column]:' + str(df.iloc[0, last_number_column]))
        if last_number_column is not None and length >= 2 and last_number_column < len(df.columns):
            print('df.iloc[1, last_number_column]:' + str(df.iloc[1, last_number_column]))
        if last_number_column is not None and length >= 3 and last_number_column < len(df.columns):
            print('df.iloc[2, last_number_column]:' + str(df.iloc[2, last_number_column]))
        print('last_num:' + str(last_num))
        current_num_row = None
        if last_number_column is not None and last_number_column < len(df.columns):
            for i in range(length):
                print('df.iloc[i, last_number_column]: ' + str(df.iloc[i, last_number_column]))
                if df.iloc[i, last_number_column] is not None and 'номер точки' in df.iloc[
                    i, last_number_column].lower().replace('\n', ' ').strip():
                    current_num_row = i
        print('current_num_row:' + str(current_num_row))
        if current_num_row is not None and length >= current_num_row + 1 and last_number_column is not None:
            print('df.iloc[current_num_row, last_number_column]:' + str(
                df.iloc[current_num_row, last_number_column]))
            print('df.iloc[current_num_row + 1, last_number_column]:' + str(
                df.iloc[current_num_row + 1, last_number_column]))
        if last_width and last_width <= len(df.columns) and \
                isinstance(last_num, str) and last_num.isdigit() and last_number_column is not None and \
                length >= 2 and width >= last_number_column and \
                (current_num_row is not None and check_table_numbers_join(current_num_row + 1, last_number_column,
                                                                          last_num, df,
                                                                          length, width) or check_table_numbers_join(0,
                                                                                                                     last_number_column,
                                                                                                                     last_num,
                                                                                                                     df,
                                                                                                                     length,
                                                                                                                     width)):
            fill_dataframe_from_pdf(last_target_cell, df, dfs, multiple_coord_sys, coordinate_systems, current_area,
                                    legend_length, True)
            appending = True
        elif (length > 0 and width > 0 and df.iloc[0, 0] is not None and ('земельный участок' in df.iloc[
            0, 0].lower() or 'государственный регистратор' in df.iloc[0, 0].lower())) or width <= 2:
            print('SKIPPED!!')
            continue
        elif found_latitude and found_longitude and target_cell:
            # Определяем границы подтаблицы
            current_area = [None]
            fill_dataframe_from_pdf(target_cell, df, dfs, multiple_coord_sys, coordinate_systems, current_area,
                                    legend_length)
            last_target_cell = target_cell
            appending = False

            '''
            if len(dfs) > 1:
                row_to_insert = dfs.pop(0).iloc[[0]]
                for i in range(len(dfs)):
                    dfs[i] = pd.concat([row_to_insert, dfs[i]], ignore_index=True)
            '''
        else:
            appending = False
            current_area = [None]

        if number_column is not None:
            if last_number_column is not None:
                print('RESELCET last_num: ' + str(last_num) + ' ' + str(df.iloc[-1, last_number_column]) + ' . ' + str(
                    df.iloc[0, last_number_column]) + ' , ' + str(last_number_column))
            if appending:
                last_num = df.iloc[-1, last_number_column]
            else:
                last_num = df.iloc[-1, number_column]
                last_number_column = number_column
            last_width = len(df.columns)
        elif not appending:
            last_num = None
            last_width = None
            last_number_column = None
    return dfs, coordinate_systems, full_text


def check_table_numbers_join(index: int, column: int, number: str, df, length: int, width: int,
                             iter_len: int = 3) -> bool:
    return any([i for i in [df.shape[0] > index + j and df.shape[1] > column and df.iloc[
        index + j, column] is not None and str.isdigit(
        df.iloc[index + j, column]) and int(number) <= int(
        df.iloc[index + j, column]) <= int(number) + 2 for j in range(iter_len)]])


def extract_tables_from_docx(docx_path):
    """Извлекает все таблицы из docx и возвращает список DataFrame"""
    try:
        doc = Document(docx_path)
        tables_data = []

        for table in doc.tables:
            data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                data.append(row_data)

            # Преобразуем в DataFrame
            df = pd.DataFrame(data[1:], columns=data[0]) if len(data) > 1 else pd.DataFrame(data)
            tables_data.append(df)

        return tables_data
    except Exception as e:
        print(f"Ошибка при извлечении таблиц из DOCX: {e}")
        return None


def extract_coordinates_from_docx_table(table, doc):
    """Анализирует таблицу из DOCX аналогично обработке Excel-файлов"""
    # Создаем временный DataFrame из таблицы docx
    data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        data.append(row_data)
    df = pd.DataFrame(data)

    # Имитируем структуру openpyxl для совместимости
    class CellMock:
        def __init__(self, value, row, col):
            self.value = value
            self.row = row + 1  # openpyxl использует 1-индексацию
            self.column = col + 1
            self.column_letter = openpyxl.utils.get_column_letter(col + 1)

    found_longitude = found_latitude = False
    coordinate_systems = []
    multiple_coord_sys = False
    target_cell = None

    # Анализируем все ячейки таблицы
    for row_idx in range(len(df)):
        for col_idx in range(len(df.columns)):
            cell_value = df.iat[row_idx, col_idx]
            if pd.isna(cell_value) or cell_value == "":
                continue

            cell_str = str(cell_value).lower().strip()
            cell_mock = CellMock(cell_value, row_idx, col_idx)

            # Поиск систем координат
            for pattern in COORDINATE_SYSTEMS:
                coord_sys = re.search(pattern, cell_str, re.IGNORECASE)
                if coord_sys:
                    coord_sys = coord_sys.group(0)
                    if coord_sys not in coordinate_systems:
                        coordinate_systems.append(coord_sys)

            # Поиск меток координат
            for key in COORDINATE_MARKS.keys():
                for pattern in key:
                    if len(cell_str) > 1 and re.search(pattern, cell_str, re.IGNORECASE) or cell_str == pattern:
                        COORDINATE_MARKS[key] = True

            # Поиск заголовков координат
            if cell_str in ('северная', 'широта', 'x', 'n') or 'север' in cell_str or 'широт' in cell_str:
                found_latitude = True
                target_cell = cell_mock
            elif cell_str in ('восточная', 'долгота', 'y', 'e') or 'вост' in cell_str or 'долг' in cell_str:
                found_longitude = True
                target_cell = cell_mock

    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    full_text = '\n'.join(full_text)
    for pattern in COORDINATE_SYSTEMS:
        coord_sys = re.search(pattern, full_text, re.IGNORECASE)
        if coord_sys:
            coord_sys = coord_sys.group(0).lower()
            if coord_sys not in coordinate_systems:
                coordinate_systems.append(coord_sys)

    # Проверка на множественные системы координат
    if sum([1 for x in COORDINATE_MARKS.values() if x is True]) > 1 and len(coordinate_systems) > 1:
        multiple_coord_sys = True

    print('multiple_coord_sys:', multiple_coord_sys)
    print(coordinate_systems)

    # Если нашли координаты, извлекаем подтаблицу
    if found_latitude and found_longitude and target_cell:
        # Определяем границы подтаблицы
        if target_cell.column - 3 > 0:
            start_col = max(0, target_cell.column - 4)  # 0-индексация для DataFrame
        else:
            start_col = max(0, target_cell.column - 3)

        if multiple_coord_sys:
            end_col = target_cell.column + (len(coordinate_systems) - 1) * 2 - 1
        else:
            end_col = target_cell.column - 1

        start_row = target_cell.row - 1  # 0-индексация
        end_row = len(df) - 1

        # Извлекаем подтаблицу
        data = [[]]
        i = 0
        current_area = None

        for r in range(start_row, end_row + 1):
            row_data = []
            for c in range(start_col, end_col + 1):
                row_data.append(df.iat[r, c] if c < len(df.columns) else None)

            # Проверка на разделитель "участок"
            if row_data[0] and 'участок' in str(row_data[0]).lower() and current_area != row_data[0]:
                current_area = row_data[0]
                data.append([])
                i += 1
            data[i].append(row_data)

        # Преобразуем в DataFrame
        dfs = []
        for frame in data:
            if frame:  # Пропускаем пустые фреймы
                dfs.append(pd.DataFrame(frame))

        if len(dfs) > 1:
            row_to_insert = dfs.pop(0).iloc[[0]]
            for i in range(len(dfs)):
                dfs[i] = pd.concat([row_to_insert, dfs[i]], ignore_index=True)

        return dfs, coordinate_systems

    return None


def extract_coordinates_xlsx(file_path):
    # Открываем указанный xlsx-файл
    workbook = openpyxl.load_workbook(file_path)
    sheet = workbook.active  # Используем активный лист
    found_longitude = found_latitude = False
    coordinate_systems = []
    multiple_coord_sys = False
    print('++()+++')

    for row in sheet.iter_rows():
        for cell in row:
            if cell.value is not None:
                cell_str = str(cell.value).lower().strip()
                for pattern in COORDINATE_SYSTEMS:
                    coord_sys = re.search(pattern, cell_str, re.IGNORECASE)
                    if coord_sys:
                        coord_sys = coord_sys.group(0).lower()
                        if coord_sys not in coordinate_systems:
                            coordinate_systems.append(coord_sys)
                for key in COORDINATE_MARKS.keys():
                    for pattern in key:
                        if len(cell_str) > 1 and re.search(pattern, cell_str, re.IGNORECASE) or cell_str == pattern:
                            COORDINATE_MARKS[key] = True
    if sum([1 for x in COORDINATE_MARKS.values() if x is True]) > 1 and len(coordinate_systems) > 1:
        multiple_coord_sys = True
    print('multiple_coord_sys:', multiple_coord_sys)
    print(coordinate_systems)

    # Проходим по первым 20 столбцам
    for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row, min_col=1, max_col=20):
        for cell in row:
            if not cell.value:
                continue
            cell_name = cell.value.lower().strip()
            if cell_name in ('северная', 'широта', 'x', 'n') or 'север' in cell_name or 'широт' in cell_name:
                found_latitude = True
            elif cell_name in ('восточная', 'долгота', 'y', 'e') or 'вост' in cell_name or 'долг' in cell_name:
                found_longitude = True
            if found_latitude and found_longitude:
                # Если нашли строку, определяем границы подтаблицы
                if cell.column - 3 > 0:
                    start_col = max(1, cell.column - 3)
                else:
                    start_col = max(1, cell.column - 2)
                if multiple_coord_sys:
                    end_col = cell.column + (len(coordinate_systems) - 1) * 2
                else:
                    end_col = cell.column
                start_row = cell.row  # Текущая строка
                # Извлекаем подтаблицу
                data = [[]]
                i = 0
                current_area = None
                for r in sheet.iter_rows(min_row=start_row, min_col=start_col, max_col=end_col):
                    if r[0].value and 'участок' in str(r[0].value).lower() and current_area != r[0].value:
                        current_area = r[0].value
                        data.append([])
                        i += 1
                    data[i].append([cell.value for cell in r])
                # Преобразуем в DataFrame
                dfs = []
                for frame in data:
                    dfs.append(pd.DataFrame(frame))
                if len(dfs) > 1:
                    row_to_insert = dfs.pop(0).iloc[[0]]
                    for i in range(len(dfs)):
                        dfs[i] = pd.concat([row_to_insert, dfs[i]], ignore_index=True)
                return dfs, coordinate_systems


def search_coords_in_text(pdf, page_number, document, tables, text, coordinates):
    found_table = False
    points_type = None
    coords_system = None

    page_tables = pdf.pages[page_number].extract_tables()
    next_page_tables = None
    if page_number + 1 < len(pdf.pages):
        next_page_tables = pdf.pages[page_number + 1].extract_tables()
    if page_number < len(document) - 1:
        next_page = document[page_number + 1]
        text += next_page.get_text()

    if not found_table:
        points_type = re.search(r'Пункты\s+фотофиксации\.\s+Система\s+координат\S*\s+\S+\d+(?:,\s*зона\s*\d+)?\b', text,
                                re.IGNORECASE | re.MULTILINE)
        if not points_type:
            points_type = re.search(
                r'Каталог\s+координат\s+Участка\.\s+Система\s+координат\S*\s+\S+\d+(?:,\s*зона\s*\d+)?\b', text,
                re.IGNORECASE | re.MULTILINE)
            if not points_type:
                points_type = re.search(
                    r'Каталог\s+координат\s+\b(?:wgs|мск|гск)-?\d+(?:,\s*зона\s*\d+)?\b', text,
                    re.IGNORECASE | re.MULTILINE)
        if not points_type:
            points_type = re.search(
                r'Точки\s+археологических\s+раскрытий\s+[\S\s]+?\d+(?:,\s*зона\s*\d+)?\b', text,
                re.IGNORECASE | re.MULTILINE)
            if points_type:
                points_type = points_type.group(0).replace('Точки археологических раскрытий', 'Шурфы').replace('\n',
                                                                                                               ' ')
        if points_type:
            if not isinstance(points_type, str):
                points_type = points_type.group(0).replace('\n', ' ')
            coords_system = re.search(r'\b(?:wgs|мск|гск)-?\d+(?:,\s*зона\s*\d+)?\b', points_type,
                                      re.IGNORECASE | re.MULTILINE)
            print('points_type: ' + str(points_type))
            print('coords_system: ' + str(coords_system))
            if coords_system:
                coords_system = coords_system.group(0).lower().replace(' ', '').replace('-', '').replace(',',
                                                                                                         '').replace(
                    ':', '')
            print('coords_system: ' + str(coords_system))
    polygon_number = None
    if points_type or found_table:
        print('TEST TABLES !!! ' + str(coordinates))
        if page_tables:
            tables_len = len(page_tables)
            for i in range(tables_len):
                if len(page_tables[i]) > 0 and len(page_tables[i][0]) > 1:
                    if check_is_coordinate_table(page_tables[i]):
                        last_tables_joined = False
                        if i + 1 < tables_len:
                            for j in range(i + 1, tables_len):
                                print('TRY1: ' + str(page_tables[i]))
                                print('TRY2: ' + str(page_tables[j]))
                                print(check_tables_joining(page_tables[i], page_tables[j]))
                                if check_tables_joining(page_tables[i], page_tables[j]):
                                    print('CHAINED!!!!')
                                    page_tables[i] += page_tables[j]
                                    if j + 1 == tables_len:
                                        last_tables_joined = True
                                else:
                                    break
                        next_page_joined = True
                        k = 2
                        print(next_page_joined)
                        print(next_page_tables)
                        print(last_tables_joined)
                        print('ъъъыу')
                        while next_page_joined and next_page_tables:
                            print('ъъъыу1')
                            if (i + 1 == tables_len or last_tables_joined) and next_page_tables:
                                print('ъъъыу3')
                                for j in range(len(next_page_tables)):
                                    print(page_tables[i])
                                    print(next_page_tables[j])
                                    print(check_tables_joining(page_tables[i], next_page_tables[j]))
                                    print('===')
                                    if check_tables_joining(page_tables[i], next_page_tables[j]):
                                        page_tables[i] += next_page_tables[j]
                                    else:
                                        next_page_joined = False
                                        break
                                # page_tables[i] += next_page_tables[0]
                            if page_number + k < len(pdf.pages):
                                next_page_tables = pdf.pages[page_number + k].extract_tables()
                                k += 1
                if len(page_tables[i]) > 1 and len(page_tables[i][0]) > 1:
                    print(page_tables[i][0], page_tables[i][1])
                    print(page_tables[i][0][1], page_tables[i][0][1])

                    if check_is_coordinate_table(page_tables[i]):
                        if 'полигон' in ''.join(page_tables[i][0]).lower():
                            columns = ['№ полигона', '№', 'Северная широта', 'Восточная долгота']
                            polygon_number = True
                        else:
                            columns = ['№', 'Северная широта', 'Восточная долгота']
                        df_new = pd.DataFrame([row[:len(columns)] for row in page_tables[i]],
                                              columns=columns)

                        for index, row in df_new.iterrows():
                            if row is None or not row['Северная широта'] or not row['Восточная долгота']:
                                continue
                            if row['Северная широта'] and row['Восточная долгота'] and (
                                    'Северная широта' in row['Северная широта'] or 'Восточная долгота' in row[
                                'Северная широта'] or \
                                    'Северная широта' in row['Восточная долгота'] or 'Восточная долгота' in row[
                                        'Восточная долгота'] or re.search(r'\bX\b', row['Северная широта'],
                                                                          re.IGNORECASE) or re.search(r'\bY\b', row[
                                'Восточная долгота']) or 'номер' in ''.join(row)):
                                continue
                            current_points_type = points_type
                            if polygon_number:
                                polygon_number = row['№ полигона']
                                current_points_type += ' [' + str(polygon_number) + ']'
                            if current_points_type not in coordinates.keys():
                                coordinates[current_points_type] = {}
                            point_number = row['№']
                            lat = row['Северная широта'].replace(',', '.')
                            lon = row['Восточная долгота'].replace(',', '.')
                            if 'wgs84' in coords_system:
                                lat = dms_to_decimal(lat)
                                lon = dms_to_decimal(lon)
                                coordinates[current_points_type]['coordinate_system'] = 'wgs84'
                            elif str_is_float(lat) and str_is_float(lon) and coords_system:
                                if coords_system:
                                    lat, lon = convert_to_wgs84(lat, lon, coords_system)
                                    coordinates[current_points_type]['coordinate_system'] = 'wgs84'  # coords_system

                            if 'S' in row['Северная широта']:
                                lat = -lat
                            if 'W' in row['Восточная долгота']:
                                lon = -lon

                            coordinates[current_points_type][point_number] = [lat, lon]
                    else:
                        tables.append(page_tables[i])
                        continue
                else:
                    tables.append(page_tables[i])
                    continue
    else:
        for table in page_tables:
            if table:
                tables.append(table)

    pits_coordinates = re.findall(
        r'Шурф\s+№\s+\d+\**[\s\S]+?Координаты\s+шурфа\s+в\s+системе\s+WGS-\d+:*\s+[NSEW]\s*?\d+°\d+\'\d+[\.,]\d+";*\s+[NSEW]\s*?\d+°\d+\'\d+[\.,]\d+"',
        text, re.IGNORECASE)
    # r'([NS])(\d{1,2})°(\d{1,2})\'(\d{1,2}\.\d{1,2})"\s+([EW])(\d{1,3})°(\d{1,2})\'(\d{1,2}\.\d{1,2})"'
    for coord in pits_coordinates:
        pit_number = re.search(r'Шурф\s+№\s+\d+\**', coord, re.IGNORECASE).group(0)
        lat = dms_to_decimal(re.search(r'[NS]\s*?\d+°\d+\'\d+[\.,]\d+"', coord, re.IGNORECASE).group(0))
        lon = dms_to_decimal(re.search(r'[EW]\s*?\d+°\d+\'\d+[\.,]\d+"', coord, re.IGNORECASE).group(0))
        if 'Шурфы' not in coordinates:
            coordinates['Шурфы'] = {'coordinate_system': 'wgs84'}
        coordinates['Шурфы'][pit_number] = [lat, lon]
    # extract_coordinates(file, document, page_number, folder, coordinates)


def check_tables_joining(page_table: list, adjoining_table: list) -> bool:
    if adjoining_table and check_is_coordinate_table(
            adjoining_table) and len(page_table[0]) == len(adjoining_table[0]) and (
            str.isdigit(adjoining_table[0][0]) and (int(adjoining_table[0][0]) == 1 or (
            str.isdigit(page_table[-1][0]) and (int(page_table[-1][0]) <= int(
        adjoining_table[0][0]) <= int(page_table[-1][0]) + 2)))):
        return True
    return False


def check_is_coordinate_table(page_table: list) -> bool:
    if (len(page_table) > 0 and len(page_table[0]) > 1) and (
            'Северная широта' in page_table[0][1] or 'Восточная долгота' in page_table[0][
        1] or re.search(r'\bX\b', page_table[0][1], re.IGNORECASE) or re.search(r'\bY\b',
                                                                                page_table[0][
                                                                                    1]) or \
            ('°' in page_table[1][1] and '\'' in page_table[1][1] and '"' in page_table[1][
                1] and
             ('N' in page_table[1][1] or 'E' in page_table[1][1] or 'W' in page_table[1][
                 1] or 'S' in page_table[1])) or re.search(r'\b\d+[\.,]*\d*\b', page_table[1][1], re.IGNORECASE)):
        return True
    return False


def format_coordinates(results, coordinate_systems):
    coordinates = {}
    new_coordinate_systems = []
    for sys in coordinate_systems:
        coords_system = re.search(r'\b(?:wgs|мск|гск)-? ?\d+(?:,\s*зона\s*\d+)?\b', sys,
                                  re.IGNORECASE | re.MULTILINE)
        if coords_system:
            coords_system = coords_system.group(0).lower().replace(' ', '').replace('-', '').replace(',',
                                                                                                     '').replace(
                ':', '')
            new_coordinate_systems.append(coords_system)
    coordinate_systems = new_coordinate_systems
    print('coordinate_systems: ' + str(coordinate_systems))

    counter = 0
    for table in results:
        inside_counter = 0
        title = {'zone': None, '№': None, 'x': None, 'y': None}
        points_type = 'Каталог координат'
        if len(results) > 1:
            counter += 1
            points_type += ' [' + str(counter) + ']'
        print('TABLE: ' + str(table))
        for index, row in table.iterrows():
            for column, cell in row.items():
                cell_str = str(cell).lower().strip()
                if cell_str in (None, '', 'none'):
                    continue
                print(
                    str(title['№'] is not None) + ' ' + str(title['x'] is not None) + ' ' + str(title['y'] is not None))
                if title['№'] is None or title['x'] is None or title['y'] is None:
                    for key_mark in COORDINATE_MARKS.keys():
                        for i, pattern_mark in enumerate(key_mark):
                            if len(cell_str) > 1 and re.search(pattern_mark, cell_str,
                                                               re.IGNORECASE) or cell_str == pattern_mark:
                                print(str(cell_str) + ' ' + str(title['x']) + ' ' + str(title['y']) + ' ' + str(
                                    column))
                                if i == 0:
                                    title['x'] = column
                                elif i == 1:
                                    title['y'] = column
                    if re.search(r'зон|участ', cell_str, re.IGNORECASE):
                        title['zone'] = column
                    elif re.search(r'№|номер|обознач', cell_str, re.IGNORECASE):
                        title['№'] = column
                    elif title['№'] is None and title['x'] is not None and title['y'] is not None and table.columns[
                        table.columns.get_loc(title['x']) - 1] not in (None, ''):
                        title['№'] = table.columns[table.columns.get_loc(title['x']) - 1]
                else:
                    if title['zone'] is None and table.columns.get_loc(title['№']) - 1 >= 0:
                        title['zone'] = table.columns[table.columns.get_loc(title['№']) - 1]
                    if title['zone'] is not None and not (
                            pd.isna(row[title['zone']]) or row[title['zone']] == "") and len(row[title['zone']]) > 1:
                        new_name = row[title['zone']].replace('\n', ' ').replace('"', '').replace("'", '')
                        if ' [' in points_type and points_type[:points_type.rfind(' [')] != new_name:
                            points_type = new_name + ' [' + str(counter) + ']'
                        else:
                            inside_counter += 1
                            points_type = new_name + ' [' + str(counter) + '-' + str(inside_counter) + ']'
                    point_number = row[title['№']]
                    if not pd.isna(point_number) and (isinstance(point_number, float) or isinstance(point_number,
                                                                                                    str) and point_number.isdigit()):
                        point_number = int(point_number)
                    lat = lon = None
                    if row[title['x']] not in (None, ''):
                        lat = str(row[title['x']]).replace(',', '.')
                    if row[title['y']] not in (None, ''):
                        lon = str(row[title['y']]).replace(',', '.')
                    if lat in (None, '') or lon in (None, '') or point_number in (None, '') or not (
                            re.search(r'\d+°\s*\d+\'\s*\d+[\.,]\d+"|\d+[.,]+\d+', lat,
                                      re.IGNORECASE) and re.search(r'\d+°\s*\d+\'\s*\d+[\.,]\d+"|\d+[.,]+\d+', lon,
                                                                   re.IGNORECASE)):
                        continue
                    if 'wgs84' in coordinate_systems and re.search(r'\d+°\s*\d+\'\s*\d+[\.,]\d+"', lat,
                                                                   re.IGNORECASE) and re.search(
                        r'\d+°\s*\d+\'\s*\d+[\.,]\d+"', lon, re.IGNORECASE):
                        print('HERE1')
                        lat = dms_to_decimal(lat)
                        lon = dms_to_decimal(lon)
                        if ' (wgs84)' not in points_type:
                            points_type += ' (wgs84)'
                        if points_type not in coordinates.keys():
                            coordinates[points_type] = {}
                        coordinates[points_type]['coordinate_system'] = 'wgs84'
                        print('HERE2')
                        print(coordinates)
                        print(coordinates[points_type])
                    else:
                        coord_sys = None
                        for coord_sys_iter in coordinate_systems:
                            if coord_sys_iter != 'wgs84' and (
                                    (re.search(r'\d+[.,]+\d+', lat, re.IGNORECASE) and 'мск' in coord_sys_iter) or (
                                    re.search(r'\d+°\s*\d+\'\s*\d+[\.,]\d+"', lat,
                                              re.IGNORECASE) and 'гск' in coord_sys_iter)):
                                coord_sys = coord_sys_iter
                        print('coord_sys: ' + str(coord_sys))
                        if coord_sys:
                            if 'гск' in coord_sys:
                                lat = dms_to_decimal(lat)
                                lon = dms_to_decimal(lon)
                            lat, lon = convert_to_wgs84(lat, lon, coord_sys)
                            if f' ({coord_sys})' not in points_type:
                                points_type += f' ({coord_sys})'
                                if points_type not in coordinates.keys():
                                    coordinates[points_type] = {}
                                coordinates[points_type]['coordinate_system'] = 'wgs84'

                    x_cmp = str(row[title['x']]).lower()
                    y_cmp = str(row[title['y']]).lower()
                    if 's' in x_cmp or 'ю' in x_cmp:
                        lat = -lat
                    if 'w' in y_cmp or 'з' in y_cmp:
                        lon = -lon

                    if points_type not in coordinates.keys():
                        coordinates[points_type] = {}
                    coordinates[points_type][point_number] = [lat, lon]
    return coordinates


def extract_text_from_pdf(pdf_path):
    """Извлекает весь текст из PDF"""
    full_text = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text.append(text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Ошибка при извлечении текста из PDF: {e}")
        return None


def extract_tables_from_pdf(pdf_path):
    """Извлекает таблицы из PDF"""
    tables = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                # text = page.get_text()
                # if re.search(r'Выписка\s+из\s+Единого\s+государственного\s+реестра', text, re.IGNORECASE):
                #    continue
                for table in page.extract_tables():
                    if table:
                        tables.append(table)
        return tables
    except Exception as e:
        print(f"Ошибка при извлечении таблиц из PDF: {e}")
        return None


def fill_dataframe_from_pdf(target_cell, df, dfs, multiple_coord_sys, coordinate_systems, current_area, legend_length,
                            append=False):
    if legend_length >= 3:
        if target_cell.column - legend_length > 0:
            start_col = max(0, target_cell.column - legend_length - 1)
        else:
            start_col = max(0, target_cell.column - legend_length)
    else:
        if target_cell.column - 3 > 0:
            start_col = max(0, target_cell.column - 4)  # 0-индексация для DataFrame
        else:
            start_col = max(0, target_cell.column - 3)

    if multiple_coord_sys:
        end_col = target_cell.column + (len(coordinate_systems) - 1) * 2 - 1
    else:
        end_col = target_cell.column - 1

    if append:
        start_row = 0
    else:
        start_row = target_cell.row - 1  # 0-индексация
    end_row = len(df) - 1

    # Извлекаем подтаблицу
    data = [[]]
    i = 0

    for r in range(start_row, end_row + 1):
        row_data = []
        for c in range(start_col, end_col + 1):
            row_data.append(df.iat[r, c] if c < len(df.columns) else None)

        if row_data[0] and (
                'участок' in str(row_data[0]).lower() or 'зон' in str(row_data[0]).lower()) and current_area[0] != \
                row_data[0] and row_data[0] != '1':
            if current_area[0] is not None:
                data.append([])
                i += 1
            current_area[0] = row_data[0]
        data[i].append(row_data)

    for frame in data:
        if frame:
            if append:
                dfs[-1] = pd.concat([dfs[-1], pd.DataFrame(frame)], ignore_index=True)
            else:
                dfs.append(pd.DataFrame(frame))
