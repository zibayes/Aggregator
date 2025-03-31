import copy
from datetime import datetime

from docx import Document
from pathlib import Path
import tkinter as tk
from tkinter import filedialog
import os
import zipfile
from celery import shared_task
from celery_progress.backend import ProgressRecorder
import redis
import requests
import json
from .models import ObjectAccountCard
from .files_saving import delete_files_in_directory, load_raw_account_cards
from .hash import calculate_file_hash
from .coordinates_extraction import dms_to_decimal

redis_client = redis.StrictRedis(host='localhost', port=6379, db=0)


def choose_file() -> str:
    # Открываем окно выбора файла
    file_path = filedialog.askopenfilename(title="Выберите DOC или DOCX файл")
    if file_path:
        return file_path


def normalize_coordinates(coord: str) -> str:
    coord = coord.strip()
    coord = coord.replace(' 0', ' ').replace(' ', '"')
    coord = coord[1:] if coord[0] == '0' else coord
    return coord


@shared_task(bind=True)
def process_account_cards(self, account_cards_ids, user_id):
    progress_recorder = ProgressRecorder(self)
    progress_recorder.set_progress(1, 100, '')
    account_cards, pages_count = load_raw_account_cards(account_cards_ids)
    # delete_files_in_directory('uploaded_files/users/' + str(user_id), uploaded_files)
    total_processed = [0]
    folder = 'uploaded_files/'
    file_groups = {}
    print("---" + str(pages_count))
    for account_card in account_cards:
        file = {'path': account_card.source, 'origin_filename': account_card.origin_filename,
                'processed': 'False', 'pages': {'processed': '0', 'all': pages_count[account_card.source]}}
        file_groups[str(account_card.id)] = file
    progress_json = {'user_id': user_id, 'file_groups': file_groups, 'file_types': 'account_cards',
                     'time_started': datetime.now().strftime(
                         "%Y-%m-%d %H:%M:%S")}
    redis_client.set(self.request.id, json.dumps(progress_json))
    progress_recorder.set_progress(total_processed[0], sum(pages_count.values()), progress_json)
    for account_card in account_cards:
        progress_json['file_groups'][str(account_card.id)]['processed'] = 'Processing'
        if not account_card.source.endswith(('.docx', '.pdf')):
            continue
        time_on_start = datetime.now()
        extract_text_tables_and_images(account_card.source, progress_recorder, pages_count,
                                       total_processed, account_card.id, progress_json, self.request.id, time_on_start)
        progress_json['file_groups'][str(account_card.id)]['pages']['processed'] = \
            progress_json['file_groups'][str(account_card.id)]['pages']['all']
        progress_json['file_groups'][str(account_card.id)]['processed'] = 'True'
        redis_client.set(self.request.id, json.dumps(progress_json))
        progress_recorder.set_progress(total_processed[0], sum(pages_count.values()), progress_json)
    progress_json['time_ended'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    return progress_json


def extract_text_tables_and_images(file, progress_recorder, pages_count, total_processed,
                                   account_card_id, progress_json, task_id, time_on_start):
    supplement_content = {
        "address": [],
        "description": [],
    }
    coordinates = {}

    account_cards = ObjectAccountCard.objects.all()
    for account_card in account_cards:
        if account_card.source and account_card.id != account_card_id and os.path.isfile(account_card.source):
            file_hash = calculate_file_hash(file)
            open_list_hash = calculate_file_hash(account_card.source)
            if file_hash == open_list_hash:
                raise FileExistsError(
                    f"Такой файл уже загружен в систему: {progress_json['file_groups'][str(account_card_id)]['origin_filename']}")

    current_account_card = ObjectAccountCard.objects.get(id=account_card_id)

    '''
    extracted_images = []
    current_part = 0
    time_on_start = datetime.now()
    for page_number in range(len(document)):
        pages_processed = total_processed[0] + page_number
        progress_json['file_groups'][str(act_id)][source_index]['pages']['processed'] = page_number
        expected_time = ((datetime.now() - time_on_start) / (pages_processed if pages_processed > 0 else 1)) * (sum(
            pages_count.values()) - pages_processed)
        total_seconds = int(expected_time.total_seconds())
        hours, remainder = divmod(total_seconds, 3600)
        minutes, seconds = divmod(remainder, 60)
        progress_json['expected_time'] = f"{hours:02}:{minutes:02}:{seconds:02}"
        redis_client.set(task_id, json.dumps(progress_json))
        progress_recorder.set_progress(pages_processed, sum(pages_count.values()),
                                       progress_json)
    '''

    doc = Document(file)
    folder = file[:file.rfind(".")]

    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

    full_text = []
    tables = []
    nested_tables = []
    for table in doc.tables:
        table_data = []
        nested_table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
                full_text.append(cell.text.strip())
                if cell.tables:
                    nested_table = cell.tables[0]
                    for nested_row in nested_table.rows:
                        nested_table_data.append([cell.text for cell in nested_row.cells])
                    nested_tables.append(nested_table_data)
            table_data.append(row_data)
        tables.append(table_data)

    i = 0
    for table in tables:
        '''
        print(i)
        print(table)
        print(len(table))
        print('*' * 50)
        '''
        if i == 1 and len(table) == 1:
            current_account_card.name = table[0][0]
        elif i == 2 and len(table) == 1:  # and 'время создания' in table[0][0].lower()
            current_account_card.creation_time = table[0][1]
        elif i == 3 and len(table) == 1:
            current_account_card.address = table[0][0]
        elif i == 4 and len(table) == 2:
            current_account_card.object_type = table[table[1].index('+')][0]
        elif i == 5 and len(table) == 2:
            current_account_card.general_classification = table[table[1].index('+')][0]
        elif i == 6 and len(table) == 2:
            current_account_card.description = table[0][0]
        elif i == 7 and len(table) == 9:
            for row in table:
                if '+' in row:
                    current_account_card.usage = row[row.index('+') - 1]
                    break
        elif i == 8 and len(table) == 1:
            current_account_card.discovery_info = table[0][0]
        i += 1

    if not os.path.exists(folder):
        os.makedirs(folder)

    image_captions = {}
    is_first = True
    with zipfile.ZipFile(file, 'r') as zip_file:
        for file_zip in sorted(zip_file.namelist()):
            if file_zip.startswith('word/media/'):
                image_name = os.path.basename(file_zip)
                image_path = os.path.join(folder, image_name)
                with open(image_path, 'wb') as img_file:
                    img_file.write(zip_file.read(file_zip))

                image_captions[image_name] = None
                category = 'address'
                for image_name in image_captions:
                    for text in full_text:
                        for part in text.split('\n'):
                            part = part.strip()
                            if 'объект расположен' in part.lower():
                                category = 'description'
                                continue
                            if current_account_card.name in part and part not in image_captions.values():
                                if not is_first:
                                    image_captions[image_name] = part
                                    break
                                else:
                                    is_first = False
                        if image_captions[image_name]:
                            break
                supplement_content[category].append({"label": image_captions[image_name], "source": image_path})

    coordinates['Каталог координат'] = {}
    for table in nested_tables:
        for row in table:
            if 'угол поворота' in row[0].lower() or 'северная широта' in row[1].lower() or 'восточная долгота' in row[
                2].lower():
                continue
            else:
                point_number = row[0]
                lat = dms_to_decimal(normalize_coordinates(row[1]))
                lon = dms_to_decimal(normalize_coordinates(row[2]))
                coordinates['Каталог координат'][point_number] = [lat, lon]
                # coordinates['GPS координаты углов поворотов объекта'][point_number] = [lat, lon]

    current_account_card.supplement = supplement_content
    current_account_card.coordinates = coordinates
    current_account_card.is_processing = False
    current_account_card.save()


@shared_task
def error_handler_account_cards(task, exception, exception_desc):
    print(f"Задача {task.id} завершилась с ошибкой: {exception} {exception_desc}")
    progress_json = json.loads(redis_client.get(task.id))
    for account_card_id, source in progress_json['file_groups'].items():
        print(account_card_id, source)
        if source['processed'] != 'True':
            account_card = ObjectAccountCard.objects.get(id=account_card_id)
            account_card.delete()
    progress_json['time_ended'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    raise type(exception)({"error_text": str(exception), "progress_json": progress_json}) from exception
