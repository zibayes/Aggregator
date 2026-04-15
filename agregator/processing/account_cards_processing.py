import copy
import io
import json
import shutil
import os
import re
import traceback
import zipfile
from datetime import datetime
from tkinter import filedialog
import torch
import torchvision
from torchvision import transforms as T
from fuzzywuzzy import fuzz

import cv2
import fitz
import numpy as np
import pytesseract
from PIL import Image
from celery import shared_task
from docx import Document
from pytesseract import Output
import logging
from typing import Dict, List, Optional, Tuple, Any

from agregator.processing.files_saving import load_raw_account_cards
from agregator.hash import calculate_file_hash
from agregator.models import ObjectAccountCard, IdentifiedArchaeologicalHeritageSite, ArchaeologicalHeritageSite
from agregator.redis_config import redis_client
from agregator.celery_task_template import process_documents
from agregator.processing.geo_utils import calculate_polygons_area, dms_to_decimal, normalize_coordinates
from agregator.processing.batch_kml_utils import KMLParser

logger = logging.getLogger(__name__)
pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'  # 'C:/Program Files/Tesseract-OCR/tesseract.exe'

os.environ["OMP_NUM_THREADS"] = "1"
os.environ["MKL_NUM_THREADS"] = "1"

min_area = 80000  # 100000
symbol_config = r'--oem 3 --psm 3 -c tessedit_char_whitelist=+'

# Константы для модели (можно вынести в настройки)
MODEL_PATH = "account_cards_segmenter.pth"  # путь к файлу модели
DEVICE = torch.device('cuda') if torch.cuda.is_available() else torch.device('cpu')
CONFIDENCE_THRESHOLD = 0.7
NUM_CLASSES = 6

CATEGORY_ID_TO_NAME = {
    1: "text",
    2: "table",
    3: "image",
    4: "image_caption",
    5: "compile_date"
}

# Ключевые фразы для идентификации разделов (как в predictor.py)
SECTION_KEYWORDS = {
    "NAME": ["наименование объекта"],
    "CREATION_TIME": ["время создания", "возникновения"],
    "ADDRESS": ["адрес", "местонахождение"],
    "OBJECT_TYPE": ["вид объекта"],
    "GENERAL_CLASSIFICATION": ["общая видовая принадлежность"],
    "DESCRIPTION": ["общее описание объекта"],
    "USAGE": ["использование объекта"],
    "DISCOVERY_INFO": ["сведения о дате", "обнаружения объекта"],
    "COMPILER": ["составитель"],
    "COMPILE_DATE": ["дата составления"],
    "COORDINATES": ["каталог координат", "углов"],
}

# Глобальная переменная для модели
_detection_model = None


def get_detection_model():
    global _detection_model
    if _detection_model is None:
        _detection_model = load_detection_model(MODEL_PATH, NUM_CLASSES, DEVICE)
    return _detection_model


def load_detection_model(weights_path: str, num_classes: int, device: torch.device):
    """Загрузка модели Faster R-CNN"""
    model = torchvision.models.detection.fasterrcnn_resnet50_fpn(
        weights=None,
        box_detections_per_img=50,
        box_nms_thresh=0.2,
        box_score_thresh=0.001,
    )
    in_features = model.roi_heads.box_predictor.cls_score.in_features
    model.roi_heads.box_predictor = torchvision.models.detection.faster_rcnn.FastRCNNPredictor(
        in_features, num_classes
    )
    model.roi_heads.score_thresh = 0.05
    model.roi_heads.nms_thresh = 0.2
    model.rpn.nms_thresh = 0.5
    model.rpn.pre_nms_top_n_train = 2000
    model.rpn.post_nms_top_n_train = 2000
    model.roi_heads.fg_iou_thresh = 0.7
    model.roi_heads.bg_iou_thresh = 0.3

    state_dict = torch.load(weights_path, map_location=device)
    model.load_state_dict(state_dict)
    model.to(device)
    model.eval()
    return model


def detect_regions(pil_image: Image.Image, confidence_threshold: float = CONFIDENCE_THRESHOLD):
    """Детекция блоков на изображении"""
    model = get_detection_model()
    transform = T.Compose([T.ToTensor()])
    image_tensor = transform(pil_image).to(DEVICE)

    with torch.no_grad():
        prediction = model([image_tensor])[0]

    keep = prediction['scores'] > confidence_threshold
    boxes = prediction['boxes'][keep].cpu().numpy()
    labels = prediction['labels'][keep].cpu().numpy()
    scores = prediction['scores'][keep].cpu().numpy()

    regions = []
    for box, label, score in zip(boxes, labels, scores):
        class_id = int(label)
        class_name = CATEGORY_ID_TO_NAME.get(class_id, "unknown")
        x1, y1, x2, y2 = box.astype(int)
        regions.append({
            'box': [x1, y1, x2, y2],
            'class': class_name,
            'score': score
        })
    return regions


def ocr_full_page(image_rgb: np.ndarray) -> List[Dict]:
    """OCR всей страницы с возвратом слов и координат"""
    gray = cv2.cvtColor(image_rgb, cv2.COLOR_RGB2GRAY)
    _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    data = pytesseract.image_to_data(thresh, lang='rus+eng', output_type=pytesseract.Output.DICT)
    words = []
    n = len(data['text'])
    for i in range(n):
        text = data['text'][i].strip()
        if not text:
            continue
        if len(text) < 2 and not text.isdigit():
            continue
        conf = int(data['conf'][i]) if data['conf'][i] != '-1' else 100
        if conf < 30:
            continue
        words.append({
            'text': text,
            'left': data['left'][i],
            'top': data['top'][i],
            'width': data['width'][i],
            'height': data['height'][i],
            'conf': conf
        })
    return words


def find_section_headers(words: List[Dict], used_sections: set) -> List[Dict]:
    """Поиск заголовков разделов с учётом уже найденных"""
    # Группировка слов в строки
    lines = {}
    for w in words:
        y = w['top']
        key = round(y / 20) * 20
        if key not in lines:
            lines[key] = []
        lines[key].append(w)

    sorted_lines = []
    for y_key in sorted(lines.keys()):
        line_words = sorted(lines[y_key], key=lambda w: w['left'])
        line_text = ' '.join([w['text'] for w in line_words]).lower()
        x_min = min(w['left'] for w in line_words)
        y_min = min(w['top'] for w in line_words)
        x_max = max(w['left'] + w['width'] for w in line_words)
        y_max = max(w['top'] + w['height'] for w in line_words)
        sorted_lines.append({
            'text': line_text,
            'bbox': [x_min, y_min, x_max, y_max],
            'y_center': (y_min + y_max) / 2
        })

    headers = []
    found = False
    for line in sorted_lines:
        line_text = line['text']
        for section, keywords in SECTION_KEYWORDS.items():
            if section in used_sections:
                continue
            for threshold in [95, 90, 80]:
                for kw in keywords:
                    ratio = fuzz.partial_ratio(kw, line_text)
                    if ratio >= threshold:
                        headers.append({
                            'section': section,
                            'bbox': line['bbox'],
                            'text': line_text,
                            'y_center': line['y_center'],
                            'match_ratio': ratio
                        })
                        used_sections.add(section)
                        found = True
                        break
                if found:
                    break
            if found:
                break
        if len(used_sections) == len(SECTION_KEYWORDS):
            break

    # Оставляем только первый (самый верхний) заголовок для каждой секции
    unique_headers = {}
    for hdr in headers:
        sec = hdr['section']
        if sec not in unique_headers or hdr['bbox'][1] < unique_headers[sec]['bbox'][1]:
            unique_headers[sec] = hdr

    return list(unique_headers.values())


def assign_blocks_to_sections(blocks: List[Dict], headers: List[Dict], max_vertical_gap: int = 200) -> List[Dict]:
    """Привязка блоков к разделам"""
    assigned = []
    for blk in blocks:
        x1, y1, x2, y2 = blk['box']
        block_top, block_bottom = y1, y2

        # Ищем заголовки, вертикально пересекающиеся с блоком
        intersecting_headers = []
        for hdr in headers:
            h_x1, h_y1, h_x2, h_y2 = hdr['bbox']
            if not (h_y2 < block_top or h_y1 > block_bottom):
                intersecting_headers.append(hdr)

        if intersecting_headers:
            best_hdr = None
            min_dist = float('inf')
            for hdr in intersecting_headers:
                h_x1, h_y1, h_x2, h_y2 = hdr['bbox']
                if x1 >= h_x2:
                    h_dist = x1 - h_x2
                elif h_x1 >= x2:
                    h_dist = h_x1 - x2
                else:
                    h_dist = 0
                if h_dist < min_dist:
                    min_dist = h_dist
                    best_hdr = hdr
            assigned_section = best_hdr['section']
            header_text = best_hdr['text']
            distance = min_dist
        else:
            best_hdr = None
            min_dist = float('inf')
            for hdr in headers:
                h_x1, h_y1, h_x2, h_y2 = hdr['bbox']
                if h_y2 < block_top:
                    dist = block_top - h_y2
                    if dist < min_dist and dist <= max_vertical_gap:
                        min_dist = dist
                        best_hdr = hdr
            assigned_section = best_hdr['section'] if best_hdr else 'UNKNOWN'
            header_text = best_hdr['text'] if best_hdr else None
            distance = min_dist if best_hdr else None

        assigned.append({
            **blk,
            'assigned_section': assigned_section,
            'header_text': header_text,
            'distance': distance
        })
    return assigned


def extract_text_from_block(image_rgb: np.ndarray, block: Dict) -> str:
    """Извлечение текста из области блока"""
    x1, y1, x2, y2 = block['box']
    crop = image_rgb[y1:y2, x1:x2]
    if crop.size == 0:
        return ""
    # Выбираем psm в зависимости от класса
    psm = '6' if block['class'] == 'table' else '1'  # 6 - unified block of text, 1 - auto
    config = f'--oem 3 --psm {psm}'
    text = pytesseract.image_to_string(crop, lang='rus+eng', config=config).strip()
    return text


def extract_coordinates_from_table_block(image_rgb: np.ndarray, block: Dict) -> Dict[str, List[float]]:
    """Извлечение координат из блока таблицы"""
    x1, y1, x2, y2 = block['box']
    crop = image_rgb[y1:y2, x1:x2]
    if crop.size == 0:
        return {}
    # Используем существующую функцию process_all_tables_universal, но ей нужен вложенный список строк.
    # Мы можем выполнить OCR таблицы и преобразовать в структуру таблицы.
    # Упрощённо: выполняем OCR с сохранением структуры (pytesseract image_to_data)
    gray = cv2.cvtColor(crop, cv2.COLOR_RGB2GRAY)
    _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY_INV)
    data = pytesseract.image_to_data(thresh, lang='rus+eng', config='--psm 6', output_type=pytesseract.Output.DICT)
    # Группируем слова в строки и столбцы - упрощённая реализация
    # Здесь можно использовать более сложный парсинг, как в оригинальном коде.
    # Для простоты вернём пустой словарь, чтобы не усложнять.
    # В реальном проекте нужно адаптировать существующий парсер таблиц.
    return {}


def choose_file() -> str:
    # Открываем окно выбора файла
    file_path = filedialog.askopenfilename(title="Выберите DOC или DOCX файл")
    if file_path:
        return file_path


def extract_text_from_image(image: cv2.UMat, psm_conf: str) -> str:
    custom_config = f'--oem 3 --psm {psm_conf}'
    extracted_text = pytesseract.image_to_string(image, lang='rus+eng', config=custom_config)
    return extracted_text


def sort_contours_custom(contours):
    # Получаем ограничивающие прямоугольники для каждого контура
    bounding_boxes = [cv2.boundingRect(c) for c in contours]

    # Сортируем сначала по y (высота)
    sorted_contours = sorted(zip(contours, bounding_boxes), key=lambda b: b[1][1])

    # Проверяем и меняем местами контуры по заданным условиям
    for i in range(len(sorted_contours) - 1):
        c1, box1 = sorted_contours[i]
        c2, box2 = sorted_contours[i + 1]

        # Получаем координаты верхнего левого и нижнего левого углов
        top_left_1 = box1[1]  # y-координата верхнего левого угла первого контура
        bottom_left_1 = box1[1] + box1[3]  # y-координата нижнего левого угла первого контура

        top_left_2 = box2[1]  # y-координата верхнего левого угла второго контура
        bottom_left_2 = box2[1] + box2[3]  # y-координата нижнего левого угла второго контура

        # Проверяем условие для обмена местами
        if ((top_left_1 <= bottom_left_2 and top_left_1 >= top_left_2) or (
                top_left_2 <= bottom_left_1 and top_left_2 >= top_left_1)) and (box1[0] > box2[0]):
            sorted_contours[i], sorted_contours[i + 1] = sorted_contours[i + 1], sorted_contours[i]

    return [c[0] for c in sorted_contours]


def smart_detect_table_structure(table: List[List[str]]) -> Tuple[Optional[int], Optional[int], Optional[int], int]:
    """
    Умное определение структуры таблицы:
    Возвращает (индекс_точки, индекс_широты, индекс_долготы, количество_строк_заголовков)
    """
    if not table or len(table) < 2:
        return None, None, None, 0

    # Словари для поиска в разных вариациях
    point_keywords = [
        r'№', r'номер', r'точк', r'угол', r'поворот', r'обозначение',
        r'характерн', r'point', r'центр'
    ]

    lat_keywords = [
        r'широт', r'latitude', r'lat', r'с\.ш\.', r'северн', r'north',
        r'с\. ш\.', r'с.ш', r'широты', r'северной'
    ]

    lon_keywords = [
        r'долгот', r'longitude', r'lon', r'в\.д\.', r'восточн', r'east',
        r'в\. д\.', r'в.д', r'долготы', r'восточной'
    ]

    # Случай 1: Таблица с мультизаголовками (как в новых примерах)
    # Ищем строки, где есть координаты в формате градусов
    data_start_row = 0
    for i, row in enumerate(table):
        if i > 5:  # Не проверяем слишком глубоко
            break

        # Проверяем, есть ли в строке координаты в формате градусов
        has_coords = False
        for cell in row:
            if isinstance(cell, str) and re.search(r'\d+°\d+\'', cell, re.IGNORECASE | re.MULTILINE):
                has_coords = True
                break

        if has_coords:
            data_start_row = i
            break

    # Определяем заголовки (строки до data_start_row)
    header_rows = table[:data_start_row] if data_start_row > 0 else []

    # Собираем все заголовки по столбцам (объединяя multi-row headers)
    max_cols = max(len(row) for row in table[:data_start_row + 3]) if table else 0
    column_titles = [''] * max_cols

    for row in header_rows:
        for col_idx, cell in enumerate(row):
            if col_idx < max_cols and cell and str(cell).strip():
                column_titles[col_idx] += ' ' + str(cell).strip()

    # Чистим заголовки
    column_titles = [title.strip() for title in column_titles]

    # Ищем индексы по ключевым словам в заголовках
    point_idx, lat_idx, lon_idx = None, None, None

    for col_idx, title in enumerate(column_titles):
        title_lower = title.lower()

        # Проверка на номер точки
        if any(re.search(keyword, title_lower, re.IGNORECASE | re.MULTILINE) for keyword in point_keywords):
            point_idx = col_idx

        # Проверка на широту
        if any(re.search(keyword, title_lower, re.IGNORECASE | re.MULTILINE) for keyword in lat_keywords):
            lat_idx = col_idx

        # Проверка на долготу
        if any(re.search(keyword, title_lower, re.IGNORECASE | re.MULTILINE) for keyword in lon_keywords):
            lon_idx = col_idx

    # Если не нашли через заголовки, ищем по содержимому первых строк данных
    if lat_idx is None or lon_idx is None:
        # Берем первые 3 строки данных
        sample_rows = table[data_start_row:data_start_row + 3]

        if sample_rows:
            # Для каждого столбца проверяем, содержит ли он координаты
            for col_idx in range(max_cols):
                has_lat_format, has_lon_format = False, False

                for row in sample_rows:
                    if col_idx < len(row):
                        cell = str(row[col_idx]).strip()

                        # Проверяем формат координат
                        if re.search(r'\d+°\d+\'', cell, re.IGNORECASE | re.MULTILINE):
                            # Проверяем, похоже ли на широту (обычно начинается с 55, 56, 57)
                            if re.match(r'^5[0-9]', cell, re.IGNORECASE | re.MULTILINE):
                                has_lat_format = True
                            # Проверяем, похоже ли на долготу (обычно начинается с 90, 91, 95, 96)
                            elif re.match(r'^9[0-9]', cell, re.IGNORECASE | re.MULTILINE):
                                has_lon_format = True

                if has_lat_format and lat_idx is None:
                    lat_idx = col_idx
                if has_lon_format and lon_idx is None:
                    lon_idx = col_idx

    # Если все еще не определили, используем эвристику:
    # Первый столбец - точка, второй - широта, третий - долгота
    if point_idx is None and lat_idx is not None and lon_idx is not None:
        # Ищем столбец, который не широта и не долгота
        for col_idx in range(max_cols):
            if col_idx != lat_idx and col_idx != lon_idx:
                point_idx = col_idx
                break

    return point_idx, lat_idx, lon_idx, data_start_row


def is_data_row(row: List[str], point_idx: int, lat_idx: int, lon_idx: int) -> bool:
    """Проверяет, является ли строка строкой с данными"""
    if not row:
        return False

    # Проверяем наличие необходимых столбцов
    max_idx = max(point_idx or 0, lat_idx or 0, lon_idx or 0)
    if len(row) <= max_idx:
        return False

    # Проверяем, что в ячейках нет заголовочных слов
    test_cells = []
    if point_idx is not None and point_idx < len(row):
        test_cells.append(str(row[point_idx]).lower())
    if lat_idx is not None and lat_idx < len(row):
        test_cells.append(str(row[lat_idx]).lower())
    if lon_idx is not None and lon_idx < len(row):
        test_cells.append(str(row[lon_idx]).lower())

    forbidden_terms = [
        'широта', 'долгота', 'latitude', 'longitude', 'lat', 'lon',
        'угол', 'поворот', '№', 'номер', 'north', 'east', 'точк',
        'северн', 'восточн', 'с.ш.', 'в.д.', 'координат'
    ]

    for cell in test_cells:
        for term in forbidden_terms:
            if term in cell:
                return False

    # Проверяем, что в ячейках с координатами есть символ градуса
    if lat_idx is not None and lat_idx < len(row):
        lat_cell = str(row[lat_idx])
        if '°' not in lat_cell:
            return False

    if lon_idx is not None and lon_idx < len(row):
        lon_cell = str(row[lon_idx])
        if '°' not in lon_cell:
            return False

    return True


def normalize_coordinates_better(coord_str: str) -> str:
    """Нормализует строку с координатами"""
    if not coord_str or not isinstance(coord_str, str):
        return ""

    # Заменяем разные символы кавычек на стандартные
    coord_str = coord_str.strip()
    coord_str = coord_str.replace('"', '"').replace('"', '"')
    coord_str = coord_str.replace("'", "'").replace("`", "'")
    coord_str = coord_str.replace("″", '"').replace("′′", '"')

    # Заменяем запятые в десятичных долях на точки
    # Ищем паттерн: градусы'минуты,секунды"
    match = re.search(r'(\d+)°\s*(\d+)[\'′]\s*([\d,]+\.?\d*)', coord_str)
    if match:
        degrees = match.group(1)
        minutes = match.group(2)
        seconds = match.group(3).replace(',', '.')
        coord_str = f"{degrees}°{minutes}'{seconds}\""

    return coord_str


def dms_to_decimal_robust(dms_str: str) -> Optional[float]:
    """Надежное преобразование DMS в десятичные градусы"""
    try:
        if not dms_str:
            return None

        # Нормализуем строку
        dms_str = normalize_coordinates_better(dms_str)

        # Убираем лишние пробелы
        dms_str = re.sub(r'\s+', '', dms_str)

        # Парсим различные форматы
        patterns = [
            r'(-?\d+)°(\d+)[\'′](\d+\.?\d*?)["″]?([NSEW])?',
            r'(-?\d+)°(\d+)\.(\d+\.?\d*)',  # Формат 55°35.123
            r'(-?\d+)°(\d+)[\'′](\d+\.?\d*)',  # Без секунд
            r'(-?\d+)°(\d+)[\'′](\d+)[\"″](\d+\.?\d*)',  # С десятичными в секундах
        ]

        for pattern in patterns:
            match = re.search(pattern, dms_str, re.IGNORECASE)
            if match:
                degrees = float(match.group(1))
                minutes = float(match.group(2))
                seconds = float(match.group(3)) if len(match.groups()) >= 3 else 0

                decimal = degrees + minutes / 60 + seconds / 3600

                # Учитываем направление (N/S/E/W)
                if len(match.groups()) >= 4 and match.group(4):
                    direction = match.group(4).upper()
                    if direction in ['S', 'W']:
                        decimal = -decimal

                return round(decimal, 10)

        # Если не нашли паттерн, пробуем парсить как десятичные градусы
        try:
            # Убираем все нецифровые символы кроме точки и минуса
            clean = re.sub(r'[^\d\.\-]', '', dms_str)
            if clean:
                return float(clean)
        except:
            pass

        return None
    except Exception as e:
        logger.error(f"Ошибка преобразования координаты {dms_str}: {e}")
        return None


def extract_points_from_table(table: List[List[str]]) -> Dict[str, List[float]]:
    """Извлекает точки из одной таблицы"""
    points = {}

    if not table:
        return points

    # Определяем структуру таблицы
    point_idx, lat_idx, lon_idx, data_start = smart_detect_table_structure(table)

    if lat_idx is None or lon_idx is None:
        logger.warning(f"Не удалось определить столбцы с координатами в таблице")
        return points

    logger.info(f"Определена структура: точка={point_idx}, широта={lat_idx}, долгота={lon_idx}, старт={data_start}")

    # Обрабатываем строки данных
    for row_idx, row in enumerate(table[data_start:], start=data_start):
        if not is_data_row(row, point_idx, lat_idx, lon_idx):
            continue

        try:
            # Извлекаем данные
            point_key = ""
            if point_idx is not None and point_idx < len(row):
                point_key = str(row[point_idx]).strip()

            # Если нет ключа, создаем его
            if not point_key:
                point_key = f"point_{len(points) + 1}"

            lat_str = str(row[lat_idx]).strip() if lat_idx < len(row) else ""
            lon_str = str(row[lon_idx]).strip() if lon_idx < len(row) else ""

            # Проверяем на пустые значения
            if not lat_str or not lon_str:
                continue

            # Преобразуем координаты
            lat = dms_to_decimal_robust(lat_str)
            lon = dms_to_decimal_robust(lon_str)

            if lat is not None and lon is not None:
                points[point_key] = [lat, lon]
                logger.debug(f"Добавлена точка {point_key}: {lat}, {lon}")
            else:
                logger.warning(f"Не удалось преобразовать координаты: {lat_str}, {lon_str}")

        except Exception as e:
            logger.error(f"Ошибка обработки строки {row_idx}: {row}, ошибка: {e}")
            continue

    return points


def process_all_tables_universal(nested_tables: List[List[List[str]]]) -> Dict[str, Any]:
    """Основная функция обработки всех таблиц"""
    coordinates = {'Каталог координат': {}}
    all_points = {}

    logger.info(f"Начинаем обработку {len(nested_tables)} таблиц")

    for table_idx, table in enumerate(nested_tables):
        if not table:
            continue

        logger.info(f"Обработка таблицы {table_idx + 1}/{len(nested_tables)}")

        # Пытаемся извлечь точки из таблицы
        try:
            points = extract_points_from_table(table)

            if points:
                logger.info(f"  Извлечено {len(points)} точек")

                # Объединяем точки, разрешая конфликты
                for key, value in points.items():
                    if key in all_points:
                        # Если точка уже существует, добавляем суффикс
                        suffix = 1
                        while f"{key}_{suffix}" in all_points:
                            suffix += 1
                        all_points[f"{key}_{suffix}"] = value
                    else:
                        all_points[key] = value
            else:
                logger.warning(f"  Не удалось извлечь точки из таблицы {table_idx + 1}")

        except Exception as e:
            logger.error(f"Критическая ошибка при обработке таблицы {table_idx + 1}: {e}")
            continue

    # Добавляем все точки в результат
    coordinates['Каталог координат'] = all_points

    # Добавляем метаданные
    if coordinates['Каталог координат']:
        coordinates['Каталог координат']['coordinate_system'] = 'wgs84'

    logger.info(f"Обработка завершена. Всего точек: {len(all_points)}")
    return coordinates


@shared_task(bind=True)
def process_account_cards(self, account_cards_ids, user_id):
    return process_documents(self, account_cards_ids, user_id, 'account_cards', model_class=ObjectAccountCard,
                             load_function=load_raw_account_cards,
                             process_function=extract_text_tables_and_images)


def extract_text_tables_and_images(file, progress_recorder, pages_count, total_processed,
                                   account_card_id, progress_json, task_id, time_on_start):
    supplement_content = {
        "address": [],
        "description": [],
    }
    coordinates = {}

    account_cards = ObjectAccountCard.objects.all()
    for account_card in account_cards:
        if account_card.source and account_card.id != account_card_id and os.path.isfile(account_card.source):
            file_hash = calculate_file_hash(file)
            open_list_hash = calculate_file_hash(account_card.source)
            if file_hash == open_list_hash:
                raise FileExistsError(
                    f"Такой файл уже загружен в систему: {progress_json['file_groups'][str(account_card_id)]['origin_filename']}")

    current_account_card = ObjectAccountCard.objects.get(id=account_card_id)

    '''
    extracted_images = []
    current_part = 0
    time_on_start = datetime.now()
    for page_number in range(len(document)):
        pages_processed = total_processed[0] + page_number
        progress_json['file_groups'][str(act_id)][source_index]['pages']['processed'] = page_number
        expected_time = ((datetime.now() - time_on_start) / (pages_processed if pages_processed > 0 else 1)) * (sum(
            pages_count.values()) - pages_processed)
        total_seconds = int(expected_time.total_seconds())
        hours, remainder = divmod(total_seconds, 3600)
        minutes, seconds = divmod(remainder, 60)
        progress_json['expected_time'] = f"{hours:02}:{minutes:02}:{seconds:02}"
        redis_client.set(task_id, json.dumps(progress_json))
        progress_recorder.set_progress(pages_processed, sum(pages_count.values()),
                                       progress_json)
    '''

    # folder = file[:file.rfind(".")]
    folder = file[:file.rfind("/") + 1] + 'Изображения'
    if not os.path.exists(folder):
        os.makedirs(folder)

    try:
        if file.endswith(('.doc', '.docx')):
            doc = Document(file)

            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)

            full_text = []
            tables = []
            nested_tables = []
            for table in doc.tables:
                table_data = []
                nested_table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                        full_text.append(cell.text.strip())
                        if cell.tables:
                            nested_table = cell.tables[0]
                            for nested_row in nested_table.rows:
                                nested_table_data.append([cell.text for cell in nested_row.cells])
                            nested_tables.append(nested_table_data)
                    table_data.append(row_data)
                tables.append(table_data)

            i = 0
            for table in tables:
                '''
                print(i)
                print(table)
                print(len(table))
                print('*' * 50)
                '''
                if i == 1 and len(table) == 1:
                    current_account_card.name = table[0][0]
                elif i == 2 and len(table) == 1:  # and 'время создания' in table[0][0].lower()
                    current_account_card.creation_time = table[0][1]
                elif i == 3 and len(table) == 1:
                    current_account_card.address = table[0][0]
                elif i == 4 and len(table) == 2:
                    current_account_card.object_type = table[table[1].index('+')][0]
                elif i == 5 and len(table) == 2:
                    current_account_card.general_classification = table[table[1].index('+')][0]
                elif i == 6 and len(table) == 2:
                    current_account_card.description = table[0][0]
                elif i == 7 and len(table) == 9:
                    for row in table:
                        if '+' in row:
                            current_account_card.usage = row[row.index('+') - 1]
                            break
                elif i == 8 and len(table) == 1:
                    current_account_card.discovery_info = table[0][0]
                elif i == 9 and len(table) == 2:
                    current_account_card.compiler = table[0][0] + ' ' + table[0][2]
                elif i == 11 and len(table) == 1:
                    current_account_card.compile_date = ''
                    for digit in table[0]:
                        current_account_card.compile_date += digit
                i += 1

            image_captions = {}
            is_first = True
            with zipfile.ZipFile(file, 'r') as zip_file:
                for file_zip in sorted(zip_file.namelist()):
                    if file_zip.startswith('word/media/'):
                        image_name = os.path.basename(file_zip)
                        image_path = os.path.join(folder, image_name)
                        with open(image_path, 'wb') as img_file:
                            img_file.write(zip_file.read(file_zip))

                        image_captions[image_name] = None
                        category = 'address'
                        for image_name in image_captions:
                            for text in full_text:
                                for part in text.split('\n'):
                                    part = part.strip()
                                    center = re.search(
                                        r'Координат[\S ]+?центр[\S ]+?WGS-\d+\)*\s*–*—*\s*[NS]\s*\d+[°\s]+\d+[\'\s]+\d+[\.,]\d+["\s]+;*\s*[EW]\s*\d+[°\s]+\d+[\'\s]+\d+[\.,]\d+"*',
                                        part, re.IGNORECASE)
                                    if center:
                                        center = center.group(0)
                                        lat = dms_to_decimal(
                                            normalize_coordinates(
                                                re.search(r'[NS]\s*\d+[°\s]+\d+[\'\s]+\d+[\.,]\d+"*', center,
                                                          re.IGNORECASE).group(0).replace('N ', '').replace(
                                                    'S ', '').strip()))
                                        lon = dms_to_decimal(
                                            normalize_coordinates(
                                                re.search(r'[EW]\s*\d+[°\s]+\d+[\'\s]+\d+[\.,]\d+"*', center,
                                                          re.IGNORECASE).group(0).replace('E ', '').replace('W ',
                                                                                                            '').strip()))
                                        if lat is None or lon is None:
                                            continue
                                        coordinates['Центр объекта'] = {}
                                        coordinates['Центр объекта']['Центр объекта'] = [lat, lon]
                                    if 'объект расположен' in part.lower():
                                        category = 'description'
                                        continue
                                    if current_account_card.name in part and part not in image_captions.values():
                                        if not is_first:
                                            image_captions[image_name] = part
                                            break
                                        else:
                                            is_first = False
                                if image_captions[image_name]:
                                    break
                        supplement_content[category].append({"label": image_captions[image_name], "source": image_path})

            coordinates['Каталог координат'] = {}
            coordinates = process_all_tables_universal(nested_tables)
            i = 0
            logger.info(f"coordinates['Каталог координат']: {coordinates['Каталог координат']}")
            while not coordinates['Каталог координат'] and i < len(tables):
                coordinates = process_all_tables_universal(tables[i])
                i += 1
            '''
            coordinates['Каталог координат'] = {}
            for table in nested_tables:
                for row in table:
                    if len(row) >= 3:
                        if 'угол поворота' in row[0].lower() or '№' in row[0].lower() or 'северная широта' in row[
                            1].lower() or 'северной широты' in row[
                            1].lower() or 'восточная долгота' in \
                                row[
                                    2].lower() or 'восточной долготы' in \
                                row[
                                    2].lower():
                            logger.info(f"table: {table}")
                            continue
                        else:
                            point_number = row[0]
                            lat = dms_to_decimal(normalize_coordinates(row[1]))
                            lon = dms_to_decimal(normalize_coordinates(row[2]))
                            if lat is None or lon is None:
                                continue
                            coordinates['Каталог координат'][point_number] = [lat, lon]
                            coordinates['Каталог координат']['coordinate_system'] = 'wgs84'
                            # coordinates['GPS координаты углов поворотов объекта'][point_number] = [lat, lon]
            '''

            points = [x for x in list(coordinates['Каталог координат'].keys()) if
                      x not in ('coordinate_system', 'area')]
            if 'Каталог координат' in coordinates and len(points) == 4:
                if intersect(coordinates['Каталог координат'][points[0]], coordinates['Каталог координат'][points[1]],
                             coordinates['Каталог координат'][points[2]],
                             coordinates['Каталог координат'][points[3]]) or intersect(
                    coordinates['Каталог координат'][points[1]], coordinates['Каталог координат'][points[2]],
                    coordinates['Каталог координат'][points[3]], coordinates['Каталог координат'][points[0]]):
                    coordinates['Каталог координат'][points[2]], coordinates['Каталог координат'][points[3]] = \
                        coordinates['Каталог координат'][points[3]], coordinates['Каталог координат'][points[2]]
                    coordinates['Каталог координат'] = {k: coordinates['Каталог координат'][k] for k in
                                                        sorted(coordinates['Каталог координат'])}

            calculate_polygons_area(coordinates)



        elif file.endswith('.pdf'):
            doc = fitz.open(file)
            used_sections = set()  # глобально найденные разделы по документу
            image_blocks = []  # список найденных изображений для привязки подписей
            caption_blocks = []  # список подписей
            for page_number in range(len(doc)):
                page = doc.load_page(page_number)
                pix = page.get_pixmap(dpi=300)
                pil_img = Image.open(io.BytesIO(pix.tobytes("png")))
                img_rgb = np.array(pil_img)  # RGB изображение
                img_bgr = cv2.cvtColor(img_rgb, cv2.COLOR_RGB2BGR)
                # ---------- 1. OCR всей страницы ----------
                words = ocr_full_page(img_rgb)
                # ---------- 2. Поиск заголовков разделов ----------
                headers = find_section_headers(words, used_sections)
                # ---------- 3. Детекция блоков моделью ----------
                blocks = detect_regions(pil_img)
                # ---------- 4. Привязка блоков к разделам ----------
                assigned_blocks = assign_blocks_to_sections(blocks, headers)
                assigned_blocks.sort(key=lambda b: b['box'][1])

                # ---------- 5. Обработка каждого блока ----------
                for blk in assigned_blocks:
                    x1, y1, x2, y2 = blk['box']
                    class_name = blk['class']
                    section = blk['assigned_section']

                    # Вырезаем ROI
                    roi_rgb = img_rgb[y1:y2, x1:x2]
                    roi_bgr = img_bgr[y1:y2, x1:x2]

                    if class_name == 'text':
                        # Раздел уже определён привязкой к заголовку
                        if section == 'UNKNOWN':
                            continue
                        # Распознаём текст блока (можно использовать psm 6 для блока текста)
                        block_text = pytesseract.image_to_string(roi_bgr, lang='rus+eng',
                                                                 config='--oem 3 --psm 6').strip()

                        if not block_text:
                            continue

                        if section == 'NAME':
                            current_account_card.name = block_text

                        elif section == 'CREATION_TIME':
                            current_account_card.creation_time = block_text

                        elif section == 'ADDRESS':
                            current_account_card.address = block_text
                            # Попутно ищем координаты центра в адресе (как в старом коде)
                            center_match = re.search(
                                r'Координат[\S ]+?центр[\S ]+?WGS-\d+\)*\s*–*—*\s*[NS]\s*\d+[\'’"°\s]+\d+[\'’"°\s]+\d+[\.,]\d+[\'’"°\s]+;*\s*[EW]\s*\d+[\'’"°\s]+\d+[\'’"°\s]+\d+[\.,]\d+[\'’"°]*',
                                block_text, re.IGNORECASE)
                            if center_match:
                                center = center_match.group(0)
                                lat = dms_to_decimal(normalize_coordinates(
                                    re.search(r'[NS]\s*\d+[\'’"°\s]+\d+[\'’"°\s]+\d+[\.,]\d+[\'’"°]*', center,
                                              re.IGNORECASE).group(0)
                                    .replace('N ', '').replace('S ', '').strip()))
                                lon = dms_to_decimal(normalize_coordinates(
                                    re.search(r'[EW]\s*\d+[\'’"°\s]+\d+[\'’"°\s]+\d+[\.,]\d+[\'’"°]*', center,
                                              re.IGNORECASE).group(0)
                                    .replace('E ', '').replace('W ', '').strip()))
                                if lat and lon:
                                    coordinates['Центр объекта'] = {'Центр объекта': [lat, lon],
                                                                    'coordinate_system': 'wgs84'}
                            # Сохраняем изображение адреса (как в старом коде)
                            '''
                            image_path = os.path.join(folder, f"address_page{page_number}.png")
                            cv2.imwrite(image_path, roi_bgr)
                            supplement_content['address'].append({"label": block_text, "source": image_path})
                            '''

                        elif section == 'OBJECT_TYPE':
                            # Ищем плюсик внутри блока
                            gray = cv2.cvtColor(roi_bgr, cv2.COLOR_BGR2GRAY)
                            _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)
                            data = pytesseract.image_to_data(thresh, config='--oem 3 --psm 6', lang='rus',
                                                             output_type=Output.DICT)

                            for i in range(len(data['text'])):
                                if data['text'][i] == '+':
                                    x = data['left'][i]
                                    h, w = roi_bgr.shape[:2]
                                    third = w // 3
                                    if x <= third:
                                        current_account_card.object_type = 'Памятник'
                                    elif third < x <= third * 2:
                                        current_account_card.object_type = 'Ансамбль'
                                    elif third * 2 < x <= third * 3:
                                        current_account_card.object_type = 'Достопримечательное место'
                                    break

                        elif section == 'GENERAL_CLASSIFICATION':
                            gray = cv2.cvtColor(roi_bgr, cv2.COLOR_BGR2GRAY)
                            _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)
                            data = pytesseract.image_to_data(thresh, config=symbol_config, lang='rus',
                                                             output_type=Output.DICT)

                            for i in range(len(data['text'])):
                                if data['text'][i] == '+':
                                    x = data['left'][i]
                                    h, w = roi_bgr.shape[:2]
                                    fourth = w // 4
                                    if x <= fourth:
                                        current_account_card.general_classification = 'Памятник археологии'
                                    elif fourth < x <= fourth * 2:
                                        current_account_card.general_classification = 'Памятник истории'
                                    elif fourth * 2 < x <= fourth * 3:
                                        current_account_card.general_classification = 'Памятник градостроительства и архитектуры'
                                    elif fourth * 3 < x <= fourth * 4:
                                        current_account_card.general_classification = 'Памятник монументального искусства'
                                    break

                        elif section == 'DESCRIPTION':
                            current_account_card.description = block_text

                        elif section == 'USAGE':
                            gray = cv2.cvtColor(roi_bgr, cv2.COLOR_BGR2GRAY)
                            _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)
                            data = pytesseract.image_to_data(thresh, config='--oem 3 --psm 6', lang='rus',
                                                             output_type=Output.DICT)

                            for i in range(len(data['text'])):
                                if data['text'][i] == '+':
                                    x = data['left'][i]
                                    y = data['top'][i]
                                    h, w = roi_bgr.shape[:2]
                                    x_second = w // 2
                                    y_tenth = h // 10
                                    if x <= x_second:
                                        if y_tenth < y <= y_tenth * 2:
                                            current_account_card.usage = 'Музеи, архивы, библиотеки'
                                        elif y_tenth * 2 < y <= y_tenth * 3:
                                            current_account_card.usage = 'Организации науки и образования'
                                        elif y_tenth * 3 < y <= y_tenth * 4:
                                            current_account_card.usage = 'Театрально-зрелищные организации'
                                        elif y_tenth * 4 < y <= y_tenth * 5:
                                            current_account_card.usage = 'Органы власти и управления'
                                        elif y_tenth * 5 < y <= y_tenth * 6:
                                            current_account_card.usage = 'Воинские части'
                                        elif y_tenth * 6 < y <= y_tenth * 7:
                                            current_account_card.usage = 'Религиозные организации'
                                        elif y_tenth * 7 < y <= y_tenth * 8:
                                            current_account_card.usage = 'Организации здравоохранения'
                                        elif y_tenth * 8 < y <= y_tenth * 9:
                                            current_account_card.usage = 'Организации транспорта'
                                        elif y_tenth * 9 < y <= y_tenth * 10:
                                            current_account_card.usage = 'Производственные организации'
                                    else:
                                        if y_tenth < y <= y_tenth * 2:
                                            current_account_card.usage = 'Организации торговли'
                                        elif y_tenth * 2 < y <= y_tenth * 3:
                                            current_account_card.usage = 'Организации общественного питания'
                                        elif y_tenth * 3 < y <= y_tenth * 4:
                                            current_account_card.usage = 'Гостиницы, отели'
                                        elif y_tenth * 4 < y <= y_tenth * 5:
                                            current_account_card.usage = 'Офисные помещения'
                                        elif y_tenth * 5 < y <= y_tenth * 6:
                                            current_account_card.usage = 'Жилье'
                                        elif y_tenth * 6 < y <= y_tenth * 7:
                                            current_account_card.usage = 'Парки, сады'
                                        elif y_tenth * 7 < y <= y_tenth * 8:
                                            current_account_card.usage = 'Некрополи, захоронения'
                                        elif y_tenth * 8 < y <= y_tenth * 9:
                                            current_account_card.usage = 'Не используется'
                                        elif y_tenth * 9 < y <= y_tenth * 10:
                                            current_account_card.usage = 'Иное'
                                    break

                        elif section == 'DISCOVERY_INFO':
                            current_account_card.discovery_info = block_text
                        elif section == 'COMPILER':
                            current_account_card.compiler = block_text
                        # COMPILE_DATE обрабатывается отдельным классом, но если попал сюда:
                        elif section == 'COMPILE_DATE':
                            current_account_card.compile_date = block_text


                    elif class_name == 'table':
                        # Обработка таблицы координат
                        coords = extract_coordinates_from_table_image(roi_bgr)
                        if coords:
                            if 'Каталог координат' not in coordinates:
                                coordinates['Каталог координат'] = {}
                            coordinates['Каталог координат'].update(coords)
                            coordinates['Каталог координат']['coordinate_system'] = 'wgs84'


                    elif class_name == 'image':
                        # Сохраняем изображение
                        image_filename = f"page{page_number}_{x1}_{y1}.png"
                        image_path = os.path.join(folder, image_filename)
                        cv2.imwrite(image_path, roi_bgr)
                        image_blocks.append({
                            'box': [x1, y1, x2, y2],
                            'path': image_path,
                            'page': page_number
                        })
                    elif class_name == 'image_caption':
                        caption_text = pytesseract.image_to_string(roi_bgr, lang='rus+eng',
                                                                   config='--oem 3 --psm 6').strip()
                        caption_blocks.append({
                            'box': [x1, y1, x2, y2],
                            'text': caption_text,
                            'page': page_number
                        })
                    elif class_name == 'compile_date':
                        text = pytesseract.image_to_string(roi_bgr, lang='rus+eng', config='--oem 3 --psm 6').strip()
                        current_account_card.compile_date = text

                MAX_CAPTION_DIST = 400
                # Фильтруем блоки этой страницы по классам
                page_images = [blk for blk in assigned_blocks if blk['class'] == 'image']
                page_captions = [blk for blk in assigned_blocks if blk['class'] == 'image_caption']

                for img_blk in page_images:
                    x1, y1, x2, y2 = img_blk['box']
                    best_cap = None
                    min_dist = float('inf')

                    # Ищем ближайшую подпись снизу
                    for cap_blk in page_captions:
                        cx1, cy1, cx2, cy2 = cap_blk['box']
                        if cy1 > y2:
                            # Горизонтальное пересечение не менее 30% ширины изображения
                            overlap = max(0, min(x2, cx2) - max(x1, cx1))
                            img_width = x2 - x1
                            if overlap < 0.3 * img_width:
                                continue
                            dist = cy1 - y2
                            if dist < min_dist and dist < MAX_CAPTION_DIST:
                                min_dist = dist
                                best_cap = cap_blk

                    # Получаем текст подписи
                    label = ""
                    if best_cap:
                        # OCR блока подписи
                        cap_roi = img_bgr[best_cap['box'][1]:best_cap['box'][3],
                                  best_cap['box'][0]:best_cap['box'][2]]
                        label = pytesseract.image_to_string(cap_roi, lang='rus+eng', config='--psm 6').strip()
                    else:
                        # Fallback: OCR области под изображением (100 пикселей вниз)
                        fallback_roi = img_bgr[y2:y2 + 100, x1:x2]
                        if fallback_roi.size > 0:
                            label = pytesseract.image_to_string(fallback_roi, lang='rus+eng',
                                                                config='--psm 7').strip()

                    # Сохраняем изображение
                    img_path = os.path.join(folder, f"page{page_number}_{x1}_{y1}.png")
                    cv2.imwrite(img_path, img_bgr[y1:y2, x1:x2])
                    supplement_content['description'].append({"label": label, "source": img_path})

            # Вычисление площади полигона
            if 'Каталог координат' in coordinates:
                calculate_polygons_area(coordinates)


    except Exception:
        logger.info(f"ACCOUNT CARDS FATAL ERROR")
        logger.info(traceback.format_exc())

    kml_path = KMLParser.find_kml_for_pdf(file, True)
    if kml_path:
        logger.info(f"📌 Найден KML файл: {kml_path}")

        kml_coordinates = {}
        try:
            if isinstance(kml_path, list):
                for path in kml_path:
                    kml_coordinates.update(KMLParser.parse_kml_file(path))
            else:
                kml_coordinates = KMLParser.parse_kml_file(kml_path)
        except Exception as e:
            traceback.print_exc()
            logger.warning(f"❌ Не удалось извлечь координаты из KML: {e}")

        if kml_coordinates:
            coordinates = kml_coordinates
            logger.info("✅ Координаты успешно заменены на достоверные из KML")

            total_objects = sum(len(category_objects) for category_objects in kml_coordinates.values())
            logger.info(f"📊 Извлечено {total_objects} объектов в {len(kml_coordinates)} категориях")

        else:
            logger.warning("❌ Не удалось извлечь координаты из KML")
    else:
        logger.info("ℹ️ KML файл не найден, используем координаты из PDF")

    current_account_card.supplement = supplement_content
    current_account_card.coordinates = coordinates
    current_account_card.is_processing = False
    current_account_card.save()
    if current_account_card.name:
        connect_account_card_to_heritage(current_account_card.name, progress_json)


def ccw(A, B, C):
    return (C[0] - A[0]) * (B[1] - A[1]) > (B[0] - A[0]) * (C[1] - A[1])


def intersect(A, B, C, D):
    return ccw(A, C, D) != ccw(B, C, D) and ccw(A, B, C) != ccw(A, B, D)


@shared_task
def error_handler_account_cards(task, exception, exception_desc):
    print(f"Задача {task.id} завершилась с ошибкой: {exception} {exception_desc}")
    progress_json = redis_client.get(task.id)
    if progress_json is None:
        progress_json = redis_client.get('celery-task-meta-' + str(task.id))
    progress_json = json.loads(progress_json)
    for account_card_id, source in progress_json['file_groups'].items():
        print(account_card_id, source)
        if source['processed'] != 'True':
            account_card = ObjectAccountCard.objects.get(id=account_card_id)
            account_card.delete()
    progress_json['time_ended'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    raise type(exception)({"error_text": str(exception), "progress_json": progress_json}) from exception


def connect_account_card_to_heritage(object_name, progress_json=None):
    account_card = ObjectAccountCard.objects.filter(name=object_name)
    heritage = IdentifiedArchaeologicalHeritageSite.objects.filter(name=object_name)
    if not heritage:
        heritage = ArchaeologicalHeritageSite.objects.filter(doc_name=object_name)
    if account_card and heritage:
        account_card = account_card[0]
        heritage = heritage[0]

        folder_to_move = account_card.source[:account_card.source.rfind('/')]
        destination_path = os.path.join(heritage.source, os.path.basename(folder_to_move))
        new_destination = destination_path[:destination_path.rfind('/') + 1] + 'Учётная карта'
        if os.path.exists(new_destination):
            return

        heritage.account_card_id = account_card.id
        heritage.save()

        shutil.move(folder_to_move, destination_path)
        os.rename(destination_path, new_destination)
        account_card.source = new_destination + account_card.source[account_card.source.rfind('/'):]
        account_card_supplement = copy.deepcopy(account_card.supplement_dict)
        for category, images in account_card_supplement.items():
            for i in range(len(images)):
                image_name = account_card_supplement[category][i]['source']
                account_card_supplement[category][i]['source'] = new_destination + '/Изображения' + image_name[
                                                                                                    image_name.rfind(
                                                                                                        '/'):]
        folder = account_card.source[:account_card.source.rfind('/') + 1]
        new_source = folder + account_card.origin_filename[
                              :account_card.origin_filename.rfind('.')] + account_card.source[
                                                                          account_card.source.rfind(
                                                                              '.'):]

        os.rename(account_card.source, new_source)
        account_card.source = new_source
        account_card.supplement = account_card_supplement
        account_card.save()
        if progress_json is not None:
            progress_json['file_groups'][str(account_card.id)]['path'] = account_card.source


def extract_coordinates_from_table_image(table_img: np.ndarray) -> Dict[str, List[float]]:
    """
    Извлекает координаты из изображения таблицы (BGR).
    Возвращает словарь вида { '1': [lat, lon], '2': [lat, lon], ... }
    """
    coords = {}
    gray = cv2.cvtColor(table_img, cv2.COLOR_BGR2GRAY)
    _, thresh = cv2.threshold(gray, 170, 255, cv2.THRESH_BINARY)
    data = pytesseract.image_to_data(thresh, config='--oem 3 --psm 3', lang='rus', output_type=Output.DICT)

    north_x = east_x = north_y = east_y = lat = lon = None
    max_offset = 100
    j = 0

    for i in range(len(data['text']) - 1):
        text_to_check = (data['text'][i] + ' ' + data['text'][i + 1]).lower().replace('\n', '')
        if 'северная широта' in text_to_check or 'восточная долгота' in text_to_check:
            y = data['top'][i]
            sub_img = table_img[y:, :]
            gray1 = cv2.cvtColor(sub_img, cv2.COLOR_BGR2GRAY)
            _, thresh1 = cv2.threshold(gray1, 170, 255, cv2.THRESH_BINARY)
            sub_data = pytesseract.image_to_data(thresh1, config='--oem 3 --psm 3', lang='rus', output_type=Output.DICT)
            for k in range(len(sub_data['text'])):
                txt = sub_data['text'][k].strip().lower()
                if 'северная' in txt:
                    north_x = sub_data['left'][k]
                elif 'восточная' in txt:
                    east_x = sub_data['left'][k]
                if north_x and north_x - max_offset <= sub_data['left'][k] <= north_x + max_offset:
                    lat = sub_data['text'][k]
                    north_y = sub_data['top'][k]
                    if len(lat) < 12:
                        if 11 <= len(lat) + len(sub_data['text'][k + 1]) <= 12 and north_y - max_offset <= \
                                sub_data['top'][k + 1] <= north_y + max_offset:
                            lat += sub_data['text'][k + 1]
                elif east_x and east_x - max_offset <= sub_data['left'][k] <= east_x + max_offset:
                    lon = sub_data['text'][k]
                    east_y = sub_data['top'][k]
                    if len(lon) < 12:
                        if 11 <= len(lon) + len(sub_data['text'][k + 1]) <= 12 and east_y - max_offset <= \
                                sub_data['top'][k + 1] <= east_y + max_offset:
                            lon += sub_data['text'][k + 1]
                if lat and lon and re.search(r'\d+[\'’"°\s]+\d+[\'’"°\s]+\d+[\.,]\d+[\'’"°\s]+', lat) and re.search(
                        r'\d+[\'’"°\s]+\d+[\'’"°\s]+\d+[\.,]\d+[\'’"°\s]+',
                        lon) and north_y - max_offset <= east_y <= north_y + max_offset:
                    lat_dec = dms_to_decimal(normalize_coordinates(lat.strip()))
                    lon_dec = dms_to_decimal(normalize_coordinates(lon.strip()))
                    if lat_dec and lon_dec:
                        coords[str(j + 1)] = [lat_dec, lon_dec]
                        j += 1
                    lat = lon = None
    return coords
