import json
import os
import re
from datetime import datetime
from tkinter import filedialog

import openpyxl
import pandas as pd
import pdfplumber
from celery import shared_task
from celery_progress.backend import ProgressRecorder
from docx import Document

from .coordinates_extraction import convert_to_wgs84, dms_to_decimal
from .files_saving import load_raw_commercial_offers
from .hash import calculate_file_hash
from .models import CommercialOffers
from .redis_config import redis_client

COORDINATE_SYSTEMS = [
    r'wgs.*?\d+',
    r'мск.*?\d+',
    r'гск.*?\d+',
]
COORDINATE_MARKS = {
    ('север', 'вост'): False,
    ('широт', 'долг'): False,
    ('x', 'y'): False,
    ('n', 'e'): False,
}
COORDINATE_TYPES = [
    r'[NS]+\d+°\s*\d+\'\s*\d+[\.,]\d+"|\d+°\s*\d+\'\s*\d+[\.,]\d+"[СЮ]+',
    r'[EW]+\d+°\s*\d+\'\s*\d+[\.,]\d+"|\d+°\s*\d+\'\s*\d+[\.,]\d+"[ВЗ]+',
    r'\d+°\s*\d+\'\s*\d+[\.,]\d+"',

    r'\d+[.,]+\d+',
]


def choose_file() -> str:
    # Открываем окно выбора файла
    file_path = filedialog.askopenfilename(title="Выберите DOC или DOCX файл")
    if file_path:
        return file_path


@shared_task(bind=True)
def process_commercial_offers(self, commercial_offers_ids, user_id):
    progress_recorder = ProgressRecorder(self)
    progress_recorder.set_progress(1, 100, '')
    commercial_offers, pages_count = load_raw_commercial_offers(commercial_offers_ids)
    # delete_files_in_directory('uploaded_files/users/' + str(user_id), uploaded_files)
    total_processed = [0]
    folder = 'uploaded_files/'
    file_groups = {}
    print("---" + str(pages_count))
    for commercial_offer in commercial_offers:
        file = {'path': commercial_offer.source, 'origin_filename': commercial_offer.origin_filename,
                'processed': 'False', 'pages': {'processed': '0', 'all': pages_count[commercial_offer.source]}}
        file_groups[str(commercial_offer.id)] = file
    progress_json = {'user_id': user_id, 'file_groups': file_groups, 'file_types': 'commercial_offers',
                     'time_started': datetime.now().strftime(
                         "%Y-%m-%d %H:%M:%S")}
    redis_client.set(self.request.id, json.dumps(progress_json))
    progress_recorder.set_progress(total_processed[0], sum(pages_count.values()), progress_json)
    for commercial_offer in commercial_offers:
        progress_json['file_groups'][str(commercial_offer.id)]['processed'] = 'Processing'
        if not commercial_offer.source.endswith(('.doc', '.docx', '.odt', '.xlsx', '.xls', '.pdf')):
            continue
        time_on_start = datetime.now()
        extract_coordinates(commercial_offer.source, progress_recorder, pages_count,
                            total_processed, commercial_offer.id, progress_json, self.request.id,
                            time_on_start)
        progress_json['file_groups'][str(commercial_offer.id)]['pages']['processed'] = \
            progress_json['file_groups'][str(commercial_offer.id)]['pages']['all']
        progress_json['file_groups'][str(commercial_offer.id)]['processed'] = 'True'
        redis_client.set(self.request.id, json.dumps(progress_json))
        progress_recorder.set_progress(total_processed[0], sum(pages_count.values()), progress_json)
    progress_json['time_ended'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    return progress_json


def extract_coordinates(file, progress_recorder, pages_count, total_processed,
                        commercial_offer_id, progress_json, task_id, time_on_start):
    coordinates = {}

    commercial_offers = CommercialOffers.objects.all()
    for commercial_offer in commercial_offers:
        if commercial_offer.source and commercial_offer.id != commercial_offer_id and os.path.isfile(
                commercial_offer.source):
            file_hash = calculate_file_hash(file)
            open_list_hash = calculate_file_hash(commercial_offer.source)
            if file_hash == open_list_hash:
                raise FileExistsError(
                    f"Такой файл уже загружен в систему: {progress_json['file_groups'][str(commercial_offer_id)]['origin_filename']}")

    current_commercial_offer = CommercialOffers.objects.get(id=commercial_offer_id)

    '''
    extracted_images = []
    current_part = 0
    time_on_start = datetime.now()
    for page_number in range(len(document)):
        pages_processed = total_processed[0] + page_number
        progress_json['file_groups'][str(act_id)][source_index]['pages']['processed'] = page_number
        expected_time = ((datetime.now() - time_on_start) / (pages_processed if pages_processed > 0 else 1)) * (sum(
            pages_count.values()) - pages_processed)
        total_seconds = int(expected_time.total_seconds())
        hours, remainder = divmod(total_seconds, 3600)
        minutes, seconds = divmod(remainder, 60)
        progress_json['expected_time'] = f"{hours:02}:{minutes:02}:{seconds:02}"
        redis_client.set(task_id, json.dumps(progress_json))
        progress_recorder.set_progress(pages_processed, sum(pages_count.values()),
                                       progress_json)
    '''

    folder = file[:file.rfind(".")]
    if not os.path.exists(folder):
        os.makedirs(folder)

    results = coordinate_systems = []

    file_lower = file.lower()
    if file_lower.endswith('.pdf'):
        tables = extract_tables_from_pdf(file)
        results, coordinate_systems = analyze_coordinates_in_tables_from_pdf(tables, file)
    elif file_lower.endswith(('.doc', '.docx', '.odt')):
        doc = Document(file)
        results = []
        for table in doc.tables:
            result, coordinate_system = extract_coordinates_from_docx_table(table, doc)
            results.append(result)
            for sys in coordinate_system:
                coordinate_systems.append(sys)
        results = [item for sublist in results for item in sublist]
    elif file_lower.endswith(('.xlsx', '.xls')):
        results, coordinate_systems = extract_coordinates_xlsx(file)

    if results is not None:
        print(results)
        coordinates = format_coordinates(results, coordinate_systems)

    current_commercial_offer.coordinates = coordinates
    current_commercial_offer.is_processing = False
    current_commercial_offer.save()


@shared_task
def error_handler_commercial_offers(task, exception, exception_desc):
    print(f"Задача {task.id} завершилась с ошибкой: {exception} {exception_desc}")
    progress_json = redis_client.get(task.id)
    if progress_json is None:
        progress_json = redis_client.get('celery-task-meta-' + str(task.id))
    progress_json = json.loads(progress_json)
    for account_card_id, source in progress_json['file_groups'].items():
        print(account_card_id, source)
        if source['processed'] != 'True':
            commercial_offer = CommercialOffers.objects.get(id=account_card_id)
            commercial_offer.delete()
    progress_json['time_ended'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    raise type(exception)({"error_text": str(exception), "progress_json": progress_json}) from exception


def extract_text_from_pdf(pdf_path):
    """Извлекает весь текст из PDF"""
    full_text = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text.append(text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Ошибка при извлечении текста из PDF: {e}")
        return None


def extract_tables_from_pdf(pdf_path):
    """Извлекает таблицы из PDF"""
    tables = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                for table in page.extract_tables():
                    if table:
                        tables.append(table)
        return tables
    except Exception as e:
        print(f"Ошибка при извлечении таблиц из PDF: {e}")
        return None


def fill_dataframe_from_pdf(target_cell, df, dfs, multiple_coord_sys, coordinate_systems, current_area, legend_length,
                            append=False):
    if legend_length >= 3:
        if target_cell.column - legend_length > 0:
            start_col = max(0, target_cell.column - legend_length - 1)
        else:
            start_col = max(0, target_cell.column - legend_length)
    else:
        if target_cell.column - 3 > 0:
            start_col = max(0, target_cell.column - 4)  # 0-индексация для DataFrame
        else:
            start_col = max(0, target_cell.column - 3)

    if multiple_coord_sys:
        end_col = target_cell.column + (len(coordinate_systems) - 1) * 2 - 1
    else:
        end_col = target_cell.column - 1

    if append:
        start_row = 0
    else:
        start_row = target_cell.row - 1  # 0-индексация
    end_row = len(df) - 1

    # Извлекаем подтаблицу
    data = [[]]
    i = 0

    for r in range(start_row, end_row + 1):
        row_data = []
        for c in range(start_col, end_col + 1):
            row_data.append(df.iat[r, c] if c < len(df.columns) else None)

        if row_data[0] and (
                'участок' in str(row_data[0]).lower() or 'зон' in str(row_data[0]).lower()) and current_area[0] != \
                row_data[0] and row_data[0] != '1':
            if current_area[0] is not None:
                data.append([])
                i += 1
            current_area[0] = row_data[0]
        data[i].append(row_data)

    for frame in data:
        if frame:
            if append:
                dfs[-1] = pd.concat([dfs[-1], pd.DataFrame(frame)], ignore_index=True)
            else:
                dfs.append(pd.DataFrame(frame))


def analyze_coordinates_in_tables_from_pdf(tables, file_path):
    """Анализирует таблицы на наличие координатных данных"""
    dfs = []
    full_text = extract_text_from_pdf(file_path)

    if not tables:
        return None, None

    last_num = last_width = last_target_cell = last_number_column = None
    appending = False

    coordinate_systems = []
    multiple_coord_sys = False
    for table_idx, table in enumerate(tables, 1):
        # Преобразуем таблицу в DataFrame
        df = pd.DataFrame(table)

        class CellMock:
            def __init__(self, value, row, col):
                self.value = value
                self.row = row + 1  # openpyxl использует 1-индексацию
                self.column = col + 1
                self.column_letter = openpyxl.utils.get_column_letter(col + 1)

        found_longitude = found_latitude = False
        legend_length = 0
        target_cell = number_column = None
        current_area = [None]

        # Анализируем все ячейки таблицы
        for row_idx in range(len(df)):
            for col_idx in range(len(df.columns)):
                cell_value = df.iat[row_idx, col_idx]
                if pd.isna(cell_value) or cell_value == "":
                    continue

                cell_str = str(cell_value).lower().strip()
                cell_mock = CellMock(cell_value, row_idx, col_idx)

                # Поиск систем координат
                for pattern in COORDINATE_SYSTEMS:
                    coord_sys = re.search(pattern, cell_str, re.IGNORECASE)
                    if coord_sys:
                        coord_sys = coord_sys.group(0)
                        if coord_sys not in coordinate_systems:
                            coordinate_systems.append(coord_sys)

                # Поиск меток координат
                for key in COORDINATE_MARKS.keys():
                    for pattern in key:
                        if len(cell_str) > 1 and re.search(pattern, cell_str, re.IGNORECASE) or cell_str == pattern:
                            COORDINATE_MARKS[key] = True

                # Поиск заголовков координат
                if cell_str in ('северная', 'широта', 'x', 'n') or 'север' in cell_str or 'широт' in cell_str:
                    found_latitude = True
                    legend_length += 1
                    target_cell = cell_mock
                elif cell_str in ('восточная', 'долгота', 'y', 'e') or 'вост' in cell_str or 'долг' in cell_str:
                    found_longitude = True
                    legend_length += 1
                    target_cell = cell_mock
                elif 'обознач' in cell_str or 'номер' in cell_str or '№' in cell_str:
                    legend_length += 1
                    number_column = col_idx

        for pattern in COORDINATE_SYSTEMS:
            coord_sys = re.search(pattern, full_text, re.IGNORECASE)
            if coord_sys:
                coord_sys = coord_sys.group(0).lower()
                if coord_sys not in coordinate_systems:
                    coordinate_systems.append(coord_sys)

        # Проверка на множественные системы координат
        if sum([1 for x in COORDINATE_MARKS.values() if x is True]) > 1 and len(coordinate_systems) > 1:
            multiple_coord_sys = True

        print('multiple_coord_sys: ' + str(multiple_coord_sys))
        print('coordinate_systems: ' + str(coordinate_systems))
        # Если нашли координаты, извлекаем подтаблицу
        print(str(found_latitude) + ' ' + str(found_longitude) + ' ' + str(target_cell))
        print(str(last_width) + ' ' + str(len(df.columns)))

        length, width = df.shape
        if width is not None and last_number_column is not None and width >= last_number_column and length > 1:
            print(
                'Stykovochnye nomera: ' + str(last_num) + ' ' + str(df.iloc[0, last_number_column]) + ' ' + str(df.iloc[
                                                                                                                    1, last_number_column]))
        if found_latitude and found_longitude and target_cell:
            # Определяем границы подтаблицы
            current_area = [None]
            fill_dataframe_from_pdf(target_cell, df, dfs, multiple_coord_sys, coordinate_systems, current_area,
                                    legend_length)
            last_target_cell = target_cell
            appending = False

            '''
            if len(dfs) > 1:
                row_to_insert = dfs.pop(0).iloc[[0]]
                for i in range(len(dfs)):
                    dfs[i] = pd.concat([row_to_insert, dfs[i]], ignore_index=True)
            '''
        elif last_width == len(df.columns) and \
                isinstance(last_num, str) and last_num.isdigit() and last_number_column is not None and \
                length >= 2 and width >= last_number_column and \
                (str(int(last_num) + 1) == df.iloc[0, last_number_column] or str(int(last_num) + 1) == df.iloc[
                    1, last_number_column]):
            fill_dataframe_from_pdf(last_target_cell, df, dfs, multiple_coord_sys, coordinate_systems, current_area,
                                    legend_length, True)
            appending = True
        else:
            appending = False
            current_area = [None]

        if number_column is not None:
            if last_number_column is not None:
                print('RESELCET last_num: ' + str(last_num) + ' ' + str(df.iloc[-1, last_number_column]) + ' . ' + str(
                    df.iloc[0, last_number_column]) + ' , ' + str(last_number_column))
            if appending:
                last_num = df.iloc[-1, last_number_column]
            else:
                last_num = df.iloc[-1, number_column]
                last_number_column = number_column
            last_width = len(df.columns)
        elif not appending:
            last_num = None
            last_width = None
            last_number_column = None
    return dfs, coordinate_systems


def extract_tables_from_docx(docx_path):
    """Извлекает все таблицы из docx и возвращает список DataFrame"""
    try:
        doc = Document(docx_path)
        tables_data = []

        for table in doc.tables:
            data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                data.append(row_data)

            # Преобразуем в DataFrame
            df = pd.DataFrame(data[1:], columns=data[0]) if len(data) > 1 else pd.DataFrame(data)
            tables_data.append(df)

        return tables_data
    except Exception as e:
        print(f"Ошибка при извлечении таблиц из DOCX: {e}")
        return None


def extract_coordinates_from_docx_table(table, doc):
    """Анализирует таблицу из DOCX аналогично обработке Excel-файлов"""
    # Создаем временный DataFrame из таблицы docx
    data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        data.append(row_data)
    df = pd.DataFrame(data)

    # Имитируем структуру openpyxl для совместимости
    class CellMock:
        def __init__(self, value, row, col):
            self.value = value
            self.row = row + 1  # openpyxl использует 1-индексацию
            self.column = col + 1
            self.column_letter = openpyxl.utils.get_column_letter(col + 1)

    found_longitude = found_latitude = False
    coordinate_systems = []
    multiple_coord_sys = False
    target_cell = None

    # Анализируем все ячейки таблицы
    for row_idx in range(len(df)):
        for col_idx in range(len(df.columns)):
            cell_value = df.iat[row_idx, col_idx]
            if pd.isna(cell_value) or cell_value == "":
                continue

            cell_str = str(cell_value).lower().strip()
            cell_mock = CellMock(cell_value, row_idx, col_idx)

            # Поиск систем координат
            for pattern in COORDINATE_SYSTEMS:
                coord_sys = re.search(pattern, cell_str, re.IGNORECASE)
                if coord_sys:
                    coord_sys = coord_sys.group(0)
                    if coord_sys not in coordinate_systems:
                        coordinate_systems.append(coord_sys)

            # Поиск меток координат
            for key in COORDINATE_MARKS.keys():
                for pattern in key:
                    if len(cell_str) > 1 and re.search(pattern, cell_str, re.IGNORECASE) or cell_str == pattern:
                        COORDINATE_MARKS[key] = True

            # Поиск заголовков координат
            if cell_str in ('северная', 'широта', 'x', 'n') or 'север' in cell_str or 'широт' in cell_str:
                found_latitude = True
                target_cell = cell_mock
            elif cell_str in ('восточная', 'долгота', 'y', 'e') or 'вост' in cell_str or 'долг' in cell_str:
                found_longitude = True
                target_cell = cell_mock

    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    full_text = '\n'.join(full_text)
    for pattern in COORDINATE_SYSTEMS:
        coord_sys = re.search(pattern, full_text, re.IGNORECASE)
        if coord_sys:
            coord_sys = coord_sys.group(0).lower()
            if coord_sys not in coordinate_systems:
                coordinate_systems.append(coord_sys)

    # Проверка на множественные системы координат
    if sum([1 for x in COORDINATE_MARKS.values() if x is True]) > 1 and len(coordinate_systems) > 1:
        multiple_coord_sys = True

    print('multiple_coord_sys:', multiple_coord_sys)
    print(coordinate_systems)

    # Если нашли координаты, извлекаем подтаблицу
    if found_latitude and found_longitude and target_cell:
        # Определяем границы подтаблицы
        if target_cell.column - 3 > 0:
            start_col = max(0, target_cell.column - 4)  # 0-индексация для DataFrame
        else:
            start_col = max(0, target_cell.column - 3)

        if multiple_coord_sys:
            end_col = target_cell.column + (len(coordinate_systems) - 1) * 2 - 1
        else:
            end_col = target_cell.column - 1

        start_row = target_cell.row - 1  # 0-индексация
        end_row = len(df) - 1

        # Извлекаем подтаблицу
        data = [[]]
        i = 0
        current_area = None

        for r in range(start_row, end_row + 1):
            row_data = []
            for c in range(start_col, end_col + 1):
                row_data.append(df.iat[r, c] if c < len(df.columns) else None)

            # Проверка на разделитель "участок"
            if row_data[0] and 'участок' in str(row_data[0]).lower() and current_area != row_data[0]:
                current_area = row_data[0]
                data.append([])
                i += 1
            data[i].append(row_data)

        # Преобразуем в DataFrame
        dfs = []
        for frame in data:
            if frame:  # Пропускаем пустые фреймы
                dfs.append(pd.DataFrame(frame))

        if len(dfs) > 1:
            row_to_insert = dfs.pop(0).iloc[[0]]
            for i in range(len(dfs)):
                dfs[i] = pd.concat([row_to_insert, dfs[i]], ignore_index=True)

        return dfs, coordinate_systems

    return None


def extract_coordinates_xlsx(file_path):
    # Открываем указанный xlsx-файл
    workbook = openpyxl.load_workbook(file_path)
    sheet = workbook.active  # Используем активный лист
    found_longitude = found_latitude = False
    coordinate_systems = []
    multiple_coord_sys = False
    print('++()+++')

    for row in sheet.iter_rows():
        for cell in row:
            if cell.value is not None:
                cell_str = str(cell.value).lower().strip()
                for pattern in COORDINATE_SYSTEMS:
                    coord_sys = re.search(pattern, cell_str, re.IGNORECASE)
                    if coord_sys:
                        coord_sys = coord_sys.group(0).lower()
                        if coord_sys not in coordinate_systems:
                            coordinate_systems.append(coord_sys)
                for key in COORDINATE_MARKS.keys():
                    for pattern in key:
                        if len(cell_str) > 1 and re.search(pattern, cell_str, re.IGNORECASE) or cell_str == pattern:
                            COORDINATE_MARKS[key] = True
    if sum([1 for x in COORDINATE_MARKS.values() if x is True]) > 1 and len(coordinate_systems) > 1:
        multiple_coord_sys = True
    print('multiple_coord_sys:', multiple_coord_sys)
    print(coordinate_systems)

    # Проходим по первым 20 столбцам
    for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row, min_col=1, max_col=20):
        for cell in row:
            if not cell.value:
                continue
            cell_name = cell.value.lower().strip()
            if cell_name in ('северная', 'широта', 'x', 'n') or 'север' in cell_name or 'широт' in cell_name:
                found_latitude = True
            elif cell_name in ('восточная', 'долгота', 'y', 'e') or 'вост' in cell_name or 'долг' in cell_name:
                found_longitude = True
            if found_latitude and found_longitude:
                # Если нашли строку, определяем границы подтаблицы
                if cell.column - 3 > 0:
                    start_col = max(1, cell.column - 3)
                else:
                    start_col = max(1, cell.column - 2)
                if multiple_coord_sys:
                    end_col = cell.column + (len(coordinate_systems) - 1) * 2
                else:
                    end_col = cell.column
                start_row = cell.row  # Текущая строка
                # Извлекаем подтаблицу
                data = [[]]
                i = 0
                current_area = None
                for r in sheet.iter_rows(min_row=start_row, min_col=start_col, max_col=end_col):
                    if r[0].value and 'участок' in str(r[0].value).lower() and current_area != r[0].value:
                        current_area = r[0].value
                        data.append([])
                        i += 1
                    data[i].append([cell.value for cell in r])
                # Преобразуем в DataFrame
                dfs = []
                for frame in data:
                    dfs.append(pd.DataFrame(frame))
                if len(dfs) > 1:
                    row_to_insert = dfs.pop(0).iloc[[0]]
                    for i in range(len(dfs)):
                        dfs[i] = pd.concat([row_to_insert, dfs[i]], ignore_index=True)
                return dfs, coordinate_systems


def format_coordinates(results, coordinate_systems):
    coordinates = {}
    new_coordinate_systems = []
    for sys in coordinate_systems:
        coords_system = re.search(r'\b(?:wgs|мск|гск)-? ?\d+(?:,\s*зона\s*\d+)?\b', sys,
                                  re.IGNORECASE | re.MULTILINE)
        if coords_system:
            coords_system = coords_system.group(0).lower().replace(' ', '').replace('-', '').replace(',',
                                                                                                     '').replace(
                ':', '')
            new_coordinate_systems.append(coords_system)
    coordinate_systems = new_coordinate_systems

    counter = 0
    for table in results:
        inside_counter = 0
        title = {'zone': None, '№': None, 'x': None, 'y': None}
        points_type = 'Каталог координат'
        if len(results) > 1:
            counter += 1
            points_type += ' [' + str(counter) + ']'
        print('TABLE: ' + str(table))
        for index, row in table.iterrows():
            for column, cell in row.items():
                cell_str = str(cell).lower().strip()
                if cell_str in (None, '', 'none'):
                    continue
                print(
                    str(title['№'] is not None) + ' ' + str(title['x'] is not None) + ' ' + str(title['y'] is not None))
                if title['№'] is None or title['x'] is None or title['y'] is None:
                    for key_mark in COORDINATE_MARKS.keys():
                        for i, pattern_mark in enumerate(key_mark):
                            if len(cell_str) > 1 and re.search(pattern_mark, cell_str,
                                                               re.IGNORECASE) or cell_str == pattern_mark:
                                print(str(cell_str) + ' ' + str(title['x']) + ' ' + str(title['y']) + ' ' + str(
                                    column))
                                if i == 0:
                                    title['x'] = column
                                elif i == 1:
                                    title['y'] = column
                    if re.search(r'зон|участ', cell_str, re.IGNORECASE):
                        title['zone'] = column
                    elif re.search(r'№|номер|обознач', cell_str, re.IGNORECASE):
                        title['№'] = column
                    elif title['№'] is None and title['x'] is not None and title['y'] is not None and table.columns[
                        table.columns.get_loc(title['x']) - 1] not in (None, ''):
                        title['№'] = table.columns[table.columns.get_loc(title['x']) - 1]
                else:
                    if title['zone'] is None and table.columns.get_loc(title['№']) - 1 >= 0:
                        title['zone'] = table.columns[table.columns.get_loc(title['№']) - 1]
                    if title['zone'] is not None and not (
                            pd.isna(row[title['zone']]) or row[title['zone']] == "") and len(row[title['zone']]) > 1:
                        new_name = row[title['zone']].replace('\n', ' ').replace('"', '').replace("'", '')
                        if ' [' in points_type and points_type[:points_type.rfind(' [')] != new_name:
                            points_type = new_name + ' [' + str(counter) + ']'
                        else:
                            inside_counter += 1
                            points_type = new_name + ' [' + str(counter) + '-' + str(inside_counter) + ']'
                    point_number = row[title['№']]
                    if not pd.isna(point_number) and (isinstance(point_number, float) or isinstance(point_number,
                                                                                                    str) and point_number.isdigit()):
                        point_number = int(point_number)
                    lat = lon = None
                    if row[title['x']] not in (None, ''):
                        lat = str(row[title['x']]).replace(',', '.')
                    if row[title['y']] not in (None, ''):
                        lon = str(row[title['y']]).replace(',', '.')
                    if lat in (None, '') or lon in (None, '') or point_number in (None, '') or not (
                            re.search(r'\d+°\s*\d+\'\s*\d+[\.,]\d+"|\d+[.,]+\d+', lat,
                                      re.IGNORECASE) and re.search(r'\d+°\s*\d+\'\s*\d+[\.,]\d+"|\d+[.,]+\d+', lon,
                                                                   re.IGNORECASE)):
                        continue
                    if 'wgs84' in coordinate_systems and re.search(r'\d+°\s*\d+\'\s*\d+[\.,]\d+"', lat,
                                                                   re.IGNORECASE) and re.search(
                        r'\d+°\s*\d+\'\s*\d+[\.,]\d+"', lon, re.IGNORECASE):
                        lat = dms_to_decimal(lat)
                        lon = dms_to_decimal(lon)
                        if ' (wgs84)' not in points_type:
                            points_type += ' (wgs84)'
                            if points_type not in coordinates.keys():
                                coordinates[points_type] = {}
                            coordinates[points_type]['coordinate_system'] = 'wgs84'
                    else:
                        coord_sys = None
                        for coord_sys_iter in coordinate_systems:
                            if coord_sys_iter != 'wgs84' and (
                                    (re.search(r'\d+[.,]+\d+', lat, re.IGNORECASE) and 'мск' in coord_sys_iter) or (
                                    re.search(r'\d+°\s*\d+\'\s*\d+[\.,]\d+"', lat,
                                              re.IGNORECASE) and 'гск' in coord_sys_iter)):
                                coord_sys = coord_sys_iter
                        print('coord_sys: ' + str(coord_sys))
                        if coord_sys:
                            if 'гск' in coord_sys:
                                lat = dms_to_decimal(lat)
                                lon = dms_to_decimal(lon)
                            lat, lon = convert_to_wgs84(lat, lon, coord_sys)
                            if f' ({coord_sys})' not in points_type:
                                points_type += f' ({coord_sys})'
                                if points_type not in coordinates.keys():
                                    coordinates[points_type] = {}
                                coordinates[points_type]['coordinate_system'] = 'wgs84'

                    x_cmp = str(row[title['x']]).lower()
                    y_cmp = str(row[title['y']]).lower()
                    if 's' in x_cmp or 'ю' in x_cmp:
                        lat = -lat
                    if 'w' in y_cmp or 'з' in y_cmp:
                        lon = -lon

                    if points_type not in coordinates.keys():
                        coordinates[points_type] = {}
                    coordinates[points_type][point_number] = [lat, lon]
    return coordinates
